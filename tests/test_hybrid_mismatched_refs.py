from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

import scholarref


HYBRID_CASES = [
    {
        "name": "hybrid_mismatch_a2v_1.docx",
        "mode": "a2v",
        "expect": "vancouver",
        "body": [
            "Mixed corpus evidence appears in the wild (Brown & Garcia, 2019).",
            "Prior synthesis also reports unstable formatting (Miller & Stone, 2018).",
        ],
        "refs": [
            "[1] Smith A, Jones B. Sleep and memory collisions. Journal of Sleep Experiments. 2020;12(3):101-111. https://doi.org/10.1234/sleep.2020.111",
            "Brown, C., and Garcia, D  (2019)  'Urban focus drift and cognition' , Cognitive Science Review, 4(2), pp. 33-44. http://example.org/noisy",
            "3) Miller P Stone R. Directed attention outdoors? Psychol Bull 2018;145(5):321-340 doi:10.1037/bul0000200",
        ],
    },
    {
        "name": "hybrid_mismatch_a2v_2.docx",
        "mode": "a2v",
        "expect": "vancouver",
        "body": [
            "Legacy exports often flatten punctuation (Lopez et al., 2021).",
            "Some records survive but with broken separators (O'Neil, 2017a).",
        ],
        "refs": [
            "(2). Hartig T Evans GW Jamner LD, Davis DS. Tracking restoration in natural and urban field settings Journal of Environmental Psychology 2003;23(2):109-123.",
            "Lopez, F. et al 2021 Brain fog after overload. Nature Communications, 12, 456.",
            "[99] O'Neil, P. (2017a). Fragmented reference without clear separators Journal of Attention Science 9(1) 1-9",
        ],
    },
    {
        "name": "hybrid_mismatch_v2a_1.docx",
        "mode": "v2a",
        "expect": "apa7",
        "body": [
            "Numbered placeholders can be packed tightly [1-3].",
            "Cross-check confirms the same cluster [1,2].",
        ],
        "refs": [
            "1] Williams C., Davis D. Overfit attention loops. J Ment Model. 2022;8(1):1-9.",
            "(2) Brown A & Stone R (2018). Context-switch penalties. Journal of Task Control, 7(2), 44-59.",
            "3. Taylor J et al. Chaotic reference parsing in the wild. Proc Data Text. 2020;15(4):200-215. https://doi.org/10.0000/demo.2020.15",
        ],
    },
    {
        "name": "hybrid_mismatch_v2a_2.docx",
        "mode": "v2a",
        "expect": "apa7",
        "body": [
            "Brittle numbering survives OCR errors [1].",
            "Continuation often bundles entries [1-3].",
        ],
        "refs": [
            "[01] Chen L, Moore T. Adaptive noise handling. Int J Parsing. 2016;3(1):10-20.",
            "2) Nguyen P. Robust metadata recovery under corruption. Journal of Reliable Systems. 2016;11(3):77-88. doi:10.5555/robust.2016.77.",
            "Three. Patel R., Kim S. Broken numbering but valid source. Data Journal. 2019;5(2):90-99. https://example.com/fulltext",
        ],
    },
    {
        "name": "hybrid_mismatch_h2v_1.docx",
        "mode": "h2v",
        "expect": "vancouver",
        "body": [
            "Hybrid bibliographies still appear in submissions (Khan and Li, 2014).",
            "Additional support is frequently malformed (Rios et al., 2013).",
        ],
        "refs": [
            "[7] Abdullah M, Yoon K. Cognitive drift in long tasks. Behav Sci. 2015;2(1):9-21.",
            "Khan, A. and Li, B. (2014) 'Hybrid citations under pressure', Journal of Formatting, 6(4), pp. 201-219.",
            "9. Rios C et al 2013 Testing metadata salvage. Journal of Experimental Cleanup 2(2):50-60",
        ],
    },
]


def _build_doc(path: Path, body_lines: list[str], refs: list[str]) -> None:
    doc = Document()
    doc.add_heading("Hybrid Reference Stress Fixture", level=1)
    for ln in body_lines:
        doc.add_paragraph(ln)
    doc.add_paragraph("References")
    for ref in refs:
        doc.add_paragraph(ref)
    doc.save(str(path))


def _reference_lines(doc: Document) -> list[str]:
    idx = -1
    for i, p in enumerate(doc.paragraphs):
        if scholarref.is_reference_header_text(p.text):
            idx = i
            break
    if idx < 0:
        return []
    return [p.text.strip() for p in doc.paragraphs[idx + 1 :] if p.text.strip()]


def _assert_vancouver_refs_clean(lines: list[str]) -> None:
    assert lines, "No references found in Vancouver output."
    nums: list[int] = []
    for ln in lines:
        m = re.match(r"^(\d+)\.\s+", ln)
        assert m is not None, f"Output line is not Vancouver-numbered: {ln}"
        nums.append(int(m.group(1)))
        assert "http://" not in ln and "https://" not in ln, f"Raw URL leaked into Vancouver output: {ln}"
        parsed = scholarref.parse_vancouver_reference_line(ln, target_style="apa7")
        assert parsed is not None, f"Vancouver output could not be reparsed: {ln}"
        assert parsed["authors"], f"Missing authors in Vancouver output: {ln}"
        assert parsed["year"], f"Missing year in Vancouver output: {ln}"
        assert parsed["title_part"], f"Missing title in Vancouver output: {ln}"
        assert parsed["journal"], f"Missing journal in Vancouver output: {ln}"
    assert nums == list(range(1, len(lines) + 1)), "Vancouver numbering is not contiguous from 1..N."


def _assert_apa_refs_clean(lines: list[str]) -> None:
    assert lines, "No references found in APA output."
    for ln in lines:
        assert not re.match(r"^\s*[\[(]?\d+[\])\.]", ln), f"Numeric lead marker leaked into APA output: {ln}"
        assert "http://" not in ln and "https://" not in ln, f"Raw URL leaked into APA output: {ln}"
        assert re.search(r"\((?:19|20)\d{2}[a-z]?\)", ln) or "(n.d.)" in ln, f"APA year token missing: {ln}"
        parsed = scholarref.parse_author_date_reference(ln)
        assert parsed["authors"], f"Missing authors in APA output: {ln}"
        assert parsed["year"], f"Missing year in APA output: {ln}"
        assert parsed["title"], f"Missing title in APA output: {ln}"
        assert parsed["journal"], f"Missing journal in APA output: {ln}"


def test_hybrid_mismatched_reference_normalization(tmp_path: Path) -> None:
    generated = []
    for case in HYBRID_CASES:
        src = tmp_path / case["name"]
        _build_doc(src, case["body"], case["refs"])
        generated.append(src)

        doc = Document(str(src))
        if case["mode"] in {"a2v", "h2v"}:
            stats = scholarref.convert_author_date_to_vancouver(doc, keep_uncited=True)
        else:
            stats = scholarref.convert_vancouver_to_author_date(
                doc,
                target_style="apa7",
                sort_references=True,
            )

        out = tmp_path / f"{src.stem}_out.docx"
        doc.save(str(out))
        out_doc = Document(str(out))
        out_refs = _reference_lines(out_doc)

        assert len(out_refs) == len(case["refs"]), f"Reference count changed for {case['name']}."
        assert stats.get("hybrid_normalized", 0) > 0, f"Hybrid fallback was not triggered for {case['name']}."

        if case["expect"] == "vancouver":
            _assert_vancouver_refs_clean(out_refs)
        else:
            _assert_apa_refs_clean(out_refs)

    assert len(generated) == 5
