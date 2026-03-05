#!/usr/bin/env python3
"""Generate local edge-case DOCX fixtures for ScholarRef testing."""

from __future__ import annotations

from pathlib import Path

from docx import Document


def _author_date_fixture(path: Path) -> None:
    doc = Document()
    doc.add_heading("Edge Case Author-Date Fixture", level=1)
    doc.add_paragraph(
        "Cognitive recovery has mixed findings (Bell et al., 2025; Bratman et al., 2015)."
    )
    doc.add_paragraph(
        "Bratman et al. (2015) report reduced rumination after nature exposure."
    )
    doc.add_paragraph(
        "Some reviews report moderation by dose (see Bell et al., 2025; Hartig et al., 2003)."
    )
    doc.add_paragraph(
        "Mixed parentheses should keep free text (Bell et al., 2025; no consensus)."
    )
    doc.add_paragraph(
        "Two-author formatting should restyle cleanly (Miller & Stone, 2019a; Miller & Stone, 2019b)."
    )

    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Context"
    table.cell(0, 1).text = "Evidence"
    table.cell(1, 0).text = "Simulator"
    table.cell(1, 1).text = "Effects replicated (Bratman et al., 2015)."

    doc.add_paragraph("References")
    doc.add_paragraph(
        "Bell, C.N. et al. (2025) 'The relationship between nature exposures and attention "
        "restoration, as moderated by exposure duration: A systematic review and meta-analysis', "
        "Journal of Environmental Psychology, 104, p. 102632. doi:10.1016/j.jenvp.2025.102632."
    )
    doc.add_paragraph(
        "Bratman, M. G., Jonides, J., Kaplan, S., et al. (2015). Nature experience reduces "
        "rumination and subgenual prefrontal cortex activation. Proceedings of the National "
        "Academy of Sciences, 112(28), 8567-8572. doi:10.1073/pnas.1510459112."
    )
    doc.add_paragraph(
        "Hartig, T., Evans, G. W., Jamner, L. D., Davis, D. S., & Garling, T. (2003). Tracking "
        "restoration in natural and urban field settings. Journal of Environmental Psychology, "
        "23(2), 109-123. doi:10.1016/S0272-4944(02)00109-3."
    )
    doc.add_paragraph(
        "Miller, P., & Stone, R. (2019a). Directed attention and effort regulation in laboratory tasks. "
        "Psychological Bulletin, 145(4), 301-320. doi:10.1037/bul0000199."
    )
    doc.add_paragraph(
        "Miller, P., & Stone, R. (2019b). Directed attention and effort regulation in field tasks. "
        "Psychological Bulletin, 145(5), 321-340. doi:10.1037/bul0000200."
    )
    doc.save(str(path))


def _vancouver_fixture(path: Path) -> None:
    doc = Document()
    doc.add_heading("Edge Case Vancouver Fixture", level=1)
    doc.add_paragraph("Cognitive recovery has mixed findings [1,2].")
    doc.add_paragraph("Dose-response patterns remain under review [1-5].")

    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Context"
    table.cell(0, 1).text = "Evidence"
    table.cell(1, 0).text = "Simulator"
    table.cell(1, 1).text = "Effects replicated [2]."

    doc.add_paragraph("References")
    doc.add_paragraph(
        "1. Bell CN, et al. The relationship between nature exposures and attention restoration, "
        "as moderated by exposure duration: A systematic review and meta-analysis. "
        "J Environ Psychol. 2025;104:102632. doi:10.1016/j.jenvp.2025.102632."
    )
    doc.add_paragraph(
        "2. Bratman MG, Jonides J, Kaplan S, et al. Nature experience reduces rumination and "
        "subgenual prefrontal cortex activation. Proc Natl Acad Sci U S A. 2015;112(28):8567-8572. "
        "doi:10.1073/pnas.1510459112."
    )
    doc.add_paragraph(
        "3. Hartig T, Evans GW, Jamner LD, et al. Tracking restoration in natural and urban "
        "field settings. J Environ Psychol. 2003;23(2):109-123. doi:10.1016/S0272-4944(02)00109-3."
    )
    doc.add_paragraph(
        "4. Miller P, Stone R. Directed attention and effort regulation in laboratory tasks. "
        "Psychol Bull. 2019;145(4):301-320. doi:10.1037/bul0000199."
    )
    doc.add_paragraph(
        "5. Miller P, Stone R. Directed attention and effort regulation in field tasks. "
        "Psychol Bull. 2019;145(5):321-340. doi:10.1037/bul0000200."
    )
    doc.save(str(path))


def _ambiguous_author_year_fixture(path: Path) -> None:
    doc = Document()
    doc.add_heading("Ambiguous Same-Year Fixture", level=1)
    doc.add_paragraph(
        "Ambiguous citation missing suffix should fail conversion (Miller & Stone, 2019)."
    )
    doc.add_paragraph("References")
    doc.add_paragraph(
        "Miller, P., & Stone, R. (2019). Directed attention and effort regulation in laboratory tasks. "
        "Psychological Bulletin, 145(4), 301-320. doi:10.1037/bul0000199."
    )
    doc.add_paragraph(
        "Miller, P., & Stone, R. (2019). Directed attention and effort regulation in field tasks. "
        "Psychological Bulletin, 145(5), 321-340. doi:10.1037/bul0000200."
    )
    doc.save(str(path))


def main() -> int:
    out_dir = Path(__file__).resolve().parent / "generated"
    out_dir.mkdir(parents=True, exist_ok=True)
    _author_date_fixture(out_dir / "edge_author_date.docx")
    _vancouver_fixture(out_dir / "edge_vancouver.docx")
    _ambiguous_author_year_fixture(out_dir / "edge_ambiguous_author_year.docx")
    print(f"Generated fixtures in: {out_dir}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
