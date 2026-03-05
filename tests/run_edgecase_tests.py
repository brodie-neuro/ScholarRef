#!/usr/bin/env python3
"""Run conversion checks against generated edge-case DOCX fixtures."""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

import scholarref
import verify_reference_integrity
from generate_edgecase_docs import main as generate_fixtures


def _assert(cond: bool, msg: str) -> None:
    if not cond:
        raise AssertionError(msg)


def _body_text_before_refs(doc: Document) -> str:
    idx = -1
    for i, p in enumerate(doc.paragraphs):
        if scholarref.is_reference_header_text(p.text):
            idx = i
            break
    paras = scholarref.collect_body_paragraphs_before_reference(doc, idx if idx >= 0 else None)
    return "\n".join((p.text or "") for p in paras)


def _reference_lines(doc: Document) -> list[str]:
    idx = -1
    for i, p in enumerate(doc.paragraphs):
        if scholarref.is_reference_header_text(p.text):
            idx = i
            break
    if idx < 0:
        return []
    return [p.text.strip() for p in doc.paragraphs[idx + 1 :] if p.text.strip()]


def run() -> int:
    generate_fixtures()
    root = Path(__file__).resolve().parent
    gen = root / "generated"

    src_author = gen / "edge_author_date.docx"
    src_van = gen / "edge_vancouver.docx"
    src_ambiguous = gen / "edge_ambiguous_author_year.docx"

    # 1) Harvard/APA-like author-date -> Vancouver should verify cleanly.
    out_h2v = gen / "edge_author_date_h2v.docx"
    doc = Document(str(src_author))
    scholarref.convert_author_date_to_vancouver(doc, keep_uncited=True)
    doc.save(str(out_h2v))
    rc = verify_reference_integrity.verify(
        source_path=str(src_author),
        output_path=str(out_h2v),
        profile="references-only",
    )
    _assert(rc == 0, "references-only verification failed for author-date -> Vancouver")

    body_h2v = _body_text_before_refs(Document(str(out_h2v)))
    _assert(not re.search(r"\([A-Z][^)]*,\s*(?:19|20)\d{2}[a-z]?[^)]*\)", body_h2v), "APA/Harvard citation remnants found after author-date -> Vancouver")
    refs_h2v = _reference_lines(Document(str(out_h2v)))
    _assert(all("et al.." not in r for r in refs_h2v), "Double period after et al. found in Vancouver references")
    _assert(all("‘" not in r and "’" not in r for r in refs_h2v), "Quoted titles remained in Vancouver references")
    _assert(all("?. " not in r for r in refs_h2v), "Question-mark double punctuation found in Vancouver references")
    _assert(all(not re.search(r":\s*pp?\.\s*(?:\.|$)", r) for r in refs_h2v), "Empty p./pp. placeholder found in Vancouver references")
    _assert(all(not re.search(r",\s*\d+\([^)]*\),\s*(?:pp?\.)?.*;\s*(?:19|20)\d{2}", r) for r in refs_h2v), "Harvard-style reference structure remained after h2v conversion")

    # 2) Vancouver -> Harvard should remove numeric brackets and use 'and' for two-author parenthetical.
    out_v2h = gen / "edge_vancouver_v2h.docx"
    doc = Document(str(src_van))
    scholarref.convert_vancouver_to_author_date(doc, target_style="harvard", sort_references=True)
    doc.save(str(out_v2h))
    body_v2h = _body_text_before_refs(Document(str(out_v2h)))
    _assert("[" not in body_v2h, "Numeric citations remained after Vancouver -> Harvard")
    _assert("Miller and Stone, 2019a" in body_v2h, "Expected Harvard 'and' parenthetical with year suffix not found")

    # 3) APA7 <-> Harvard restyle should switch '&' <-> 'and' in parenthetical two-author citations.
    out_a2h = gen / "edge_author_date_a2h.docx"
    doc = Document(str(src_author))
    scholarref.convert_author_date_to_author_date(doc, target_style="harvard")
    doc.save(str(out_a2h))
    body_a2h = _body_text_before_refs(Document(str(out_a2h)))
    _assert("(Miller and Stone, 2019a" in body_a2h, "APA->Harvard did not restyle '&' to 'and'")

    out_h2a = gen / "edge_author_date_h2a.docx"
    doc = Document(str(out_a2h))
    scholarref.convert_author_date_to_author_date(doc, target_style="apa7")
    doc.save(str(out_h2a))
    body_h2a = _body_text_before_refs(Document(str(out_h2a)))
    _assert("(Miller & Stone, 2019a" in body_h2a, "Harvard->APA did not restyle 'and' to '&'")

    # 4) Ambiguous same-author/year citations without suffix must fail fast.
    doc = Document(str(src_ambiguous))
    raised = False
    try:
        scholarref.convert_author_date_to_vancouver(doc, keep_uncited=True)
    except RuntimeError as exc:
        raised = "Ambiguous same-author/year citations" in str(exc)
    _assert(raised, "Ambiguous same-year citations did not fail with explicit error")

    print("Edge-case DOCX tests: PASS")
    print(f"Generated outputs: {gen}")
    return 0


if __name__ == "__main__":
    raise SystemExit(run())
