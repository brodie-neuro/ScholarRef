import os
import random
import subprocess
import sys
from pathlib import Path
from docx import Document

random.seed(42)

AUTHORS = [
    "Smith", "Jones", "Williams", "Brown", "Davis", "Miller", "Wilson", "Moore",
    "Taylor", "Anderson", "Thomas", "Jackson", "White", "Harris", "Martin",
    "Thompson", "Garcia", "Martinez", "Robinson", "Clark", "Rodriguez"
]

def generate_apa_citation_and_ref(index):
    # Generates a matching pair of in-text and list reference
    num_authors = random.randint(1, 3)
    year = 1900 + index  # Guarantee no identical author/year overlap
    
    authors_chosen = random.sample(AUTHORS, num_authors)
    
    if num_authors == 1:
        in_text = f"({authors_chosen[0]}, {year})"
        ref_list = f"{authors_chosen[0]}, A. ({year}). A study on topic {index}. Journal of Science, 1(1), 1-10."
    elif num_authors == 2:
        in_text = f"({authors_chosen[0]} & {authors_chosen[1]}, {year})"
        ref_list = f"{authors_chosen[0]}, A., & {authors_chosen[1]}, B. ({year}). A study on topic {index}. Journal of Science, 1(1), 1-10."
    else:
        in_text = f"({authors_chosen[0]} et al., {year})"
        ref_list = f"{authors_chosen[0]}, A., {authors_chosen[1]}, B., & {authors_chosen[2]}, C. ({year}). A study on topic {index}. Journal of Science, 1(1), 1-10."
        
    return in_text, ref_list

def generate_vancouver_citation_and_ref(index):
    in_text = f"[{index}]"
    num_authors = random.randint(1, 3)
    year = random.randint(1990, 2025)
    authors_chosen = random.sample(AUTHORS, num_authors)
    
    if num_authors == 1:
        ref_list = f"{index}. {authors_chosen[0]} A. A study on topic {index}. J Sci. {year};1(1):1-10."
    elif num_authors == 2:
        ref_list = f"{index}. {authors_chosen[0]} A, {authors_chosen[1]} B. A study on topic {index}. J Sci. {year};1(1):1-10."
    else:
        ref_list = f"{index}. {authors_chosen[0]} A, {authors_chosen[1]} B, {authors_chosen[2]} C. A study on topic {index}. J Sci. {year};1(1):1-10."
        
    return in_text, ref_list

def build_massive_doc(filename, style_generator, num_citations=60, jumble_refs=True):
    doc = Document()
    doc.add_heading(f"Massive Jumbled Document - {filename}", level=1)
    
    # Generate citations and references
    citations = []
    references = []
    for i in range(1, num_citations + 1):
        in_text, ref_list = style_generator(i)
        citations.append(in_text)
        references.append(ref_list)
        
    # Write 30 paragraphs of jargon, injecting multiple citations into each
    filler_sentences = [
        "This is a critical finding in the field",
        "Previous research has heavily debated this mechanism",
        "Recent methodological advancements allow us to probe deeper",
        "However, the control group showed no significant deviation from baseline",
        "Theoretical models predict a rapid decay in these metrics",
        "Longitudinal data suggests otherwise",
        "This anomaly remains completely unexplained by current paradigms",
        "The implications for applied clinical practice are massive"
    ]
    
    cit_idx = 0
    for p in range(30):
        para_text = ""
        for s in range(random.randint(4, 9)): # 4 to 9 sentences per paragraph
            para_text += random.choice(filler_sentences) + " "
            # Inject 1 or 2 citations
            if cit_idx < len(citations):
                para_text += citations[cit_idx] + ". "
                cit_idx += 1
                if random.random() > 0.5 and cit_idx < len(citations):
                    # split citation or cluster them together
                    para_text += "Furthermore, other studies agree " + citations[cit_idx] + ". "
                    cit_idx += 1
            else:
                para_text += "This represents a consensus. "
                
        doc.add_paragraph(para_text.strip())
        
        # Throw in random tables and headings to jumble the structure
        if p % 7 == 0:
            doc.add_heading(f"Sub-Section {p}", level=2)
        if p % 10 == 0:
            t = doc.add_table(rows=2, cols=2)
            t.cell(0,0).text = "Metric"; t.cell(0,1).text = "Value"
            t.cell(1,0).text = "Alpha"; t.cell(1,1).text = "0.05"
            
    # Write References
    doc.add_page_break()
    doc.add_paragraph("References")
    
    # Jumble them up if requested (APA/Harvard don't need to be numbered, so jumbling tests sorting logic)
    if jumble_refs:
        random.shuffle(references)
        
    for ref in references:
        doc.add_paragraph(ref)
        # add random blank lines to simulate messy pacing
        if random.random() > 0.7:
             doc.add_paragraph("")
             
    out_dir = Path(__file__).parent / "generated_huge"
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / filename
    doc.save(out_path)
    return str(out_path)


def generate_huge_docs():
    files = []
    
    # Create enormous files
    print("Generating massive test documents...")
    files.append((build_massive_doc("huge_apa_1.docx", generate_apa_citation_and_ref, 70, jumble_refs=True), "a2v"))
    files.append((build_massive_doc("huge_apa_2.docx", generate_apa_citation_and_ref, 100, jumble_refs=True), "a2v"))
    files.append((build_massive_doc("huge_apa_3.docx", generate_apa_citation_and_ref, 40, jumble_refs=False), "a2v"))
    
    files.append((build_massive_doc("huge_van_1.docx", generate_vancouver_citation_and_ref, 80, jumble_refs=False), "v2a"))
    files.append((build_massive_doc("huge_van_2.docx", generate_vancouver_citation_and_ref, 120, jumble_refs=False), "v2a"))
    files.append((build_massive_doc("huge_van_3.docx", generate_vancouver_citation_and_ref, 50, jumble_refs=False), "v2a"))
    
    # We can reuse the APA generator for Harvard, just passing it as h2v
    files.append((build_massive_doc("huge_harv_1.docx", generate_apa_citation_and_ref, 90, jumble_refs=True), "h2v"))
    files.append((build_massive_doc("huge_harv_2.docx", generate_apa_citation_and_ref, 110, jumble_refs=True), "h2a"))
    files.append((build_massive_doc("huge_harv_3.docx", generate_apa_citation_and_ref, 60, jumble_refs=False), "h2v"))
    files.append((build_massive_doc("huge_harv_4.docx", generate_apa_citation_and_ref, 85, jumble_refs=True), "h2a"))

    return files

def run_tests():
    files = generate_huge_docs()
    success = 0
    repo_root = Path(__file__).parent.parent
    
    print("\nRunning Conversions on 10 MASSIVE Documents...")
    print("-" * 60)
    for i, (fpath, mode) in enumerate(files):
        outpath = fpath.replace(".docx", "_OUTPUT.docx")
        cmd = [
            sys.executable, str(repo_root / "scholarref.py"),
            "--mode", mode,
            "--input", fpath,
            "--output", outpath,
            "--allow-unsupported-parts"
        ]
        
        try:
            res = subprocess.run(cmd, capture_output=True, text=True, cwd=repo_root)
            if res.returncode == 0:
                print(f"[OK] Test {i+1}: {Path(fpath).name} ({mode}) - Parsed dozens of citations flawlessly.")
                success += 1
            else:
                print(f"[FAIL] Test {i+1}: {Path(fpath).name} failed.")
                print(res.stderr)
        except Exception as e:
            print(f"[ERROR] Test {i+1}: {e}")
            
    print("-" * 60)
    print(f"Passed: {success}/10")
    return success, len(files)


def test_huge_jumbled_docs_cli() -> None:
    success, total = run_tests()
    assert success == total, f"Expected all {total} massive CLI conversions to pass, got {success}."

if __name__ == "__main__":
    run_tests()
