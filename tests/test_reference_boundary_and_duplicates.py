from __future__ import annotations

import sys
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

import scholarref


def _build_doc(paragraphs: list[str], path: Path) -> None:
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    doc.save(str(path))


def _paragraph_texts(doc: Document) -> list[str]:
    return [p.text for p in doc.paragraphs]


def _reference_lines_until_break(doc: Document) -> list[str]:
    idx = -1
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() in {"References", "Reference List"}:
            idx = i
            break
    if idx < 0:
        return []

    out: list[str] = []
    for p in doc.paragraphs[idx + 1 :]:
        text = p.text.strip()
        if not text:
            continue
        if text in {"Appendix", "Supplementary checklist"}:
            break
        out.append(text)
    return out


def test_appendix_after_references_is_preserved(tmp_path: Path) -> None:
    src = tmp_path / "appendix_after_refs.docx"
    _build_doc(
        [
            "Body cites (Smith, 2020).",
            "References",
            "Smith, A. (2020). Test title. Journal of Testing, 1(1), 1-2.",
            "Appendix",
            "Supplementary checklist",
        ],
        src,
    )

    doc = Document(str(src))
    stats = scholarref.convert_author_date_to_vancouver(doc, keep_uncited=True)
    texts = _paragraph_texts(doc)

    assert stats["reference_count"] == 1
    assert "Appendix" in texts
    assert "Supplementary checklist" in texts
    assert _reference_lines_until_break(doc) == [
        "1. Smith A. Test title. Journal of Testing. 2020;1(1):1-2."
    ]


def test_exact_duplicate_author_date_references_are_collapsed(tmp_path: Path) -> None:
    src = tmp_path / "duplicate_refs.docx"
    ref_line = "Smith, A. (2020). Same paper. Journal of Testing, 1(1), 1-2."
    _build_doc(
        [
            "Body cites (Smith, 2020).",
            "References",
            ref_line,
            ref_line,
        ],
        src,
    )

    doc = Document(str(src))
    stats = scholarref.convert_author_date_to_vancouver(doc, keep_uncited=True)
    texts = _paragraph_texts(doc)

    assert stats["duplicate_refs_collapsed"] == 1
    assert stats["reference_count"] == 1
    assert texts[0] == "Body cites [1]."
    assert texts[2:] == ["1. Smith A. Same paper. Journal of Testing. 2020;1(1):1-2."]
