import os
import subprocess
import sys
from pathlib import Path
from docx import Document

def save_doc(doc, filename):
    out_dir = Path(__file__).parent / "generated"
    out_dir.mkdir(parents=True, exist_ok=True)
    doc.save(out_dir / filename)
    return str(out_dir / filename)

def generate_docs():
    files = []
    
    # 1. Standard APA 7
    doc = Document()
    doc.add_heading("1. Standard APA", level=1)
    doc.add_paragraph("This is a standard APA citation (Smith & Jones, 2020).")
    doc.add_paragraph("References")
    doc.add_paragraph("Smith, A., & Jones, B. (2020). A standard paper. Journal of Testing, 1(1), 1-10.")
    files.append((save_doc(doc, "doc1_apa_standard.docx"), "a2v"))

    # 2. APA 7 Narrative
    doc = Document()
    doc.add_heading("2. Narrative APA", level=1)
    doc.add_paragraph("As Smith and Jones (2020) argued, narrative is hard.")
    doc.add_paragraph("References")
    doc.add_paragraph("Smith, A., & Jones, B. (2020). Narrative papers. Journal of Narrative, 1(1), 1-10.")
    files.append((save_doc(doc, "doc2_apa_narrative.docx"), "a2v"))

    # 3. APA 7 Multiple Authors / et al
    doc = Document()
    doc.add_heading("3. Et Al APA", level=1)
    doc.add_paragraph("This has many authors (Williams et al., 2021).")
    doc.add_paragraph("References")
    doc.add_paragraph("Williams, C., Davis, D., Evans, E., & Ford, F. (2021). Many authors. Journal of Many.")
    files.append((save_doc(doc, "doc3_apa_etal.docx"), "a2v"))

    # 4. APA 7 Messy Split Citations
    doc = Document()
    doc.add_heading("4. Messy Split APA", level=1)
    doc.add_paragraph("Here is a split citation with extra text (see Smith & Jones, 2020; also Williams et al., 2021, for review).")
    doc.add_paragraph("References")
    doc.add_paragraph("Smith, A., & Jones, B. (2020). A standard paper. Journal of Testing, 1(1), 1-10.")
    doc.add_paragraph("Williams, C., Davis, D., Evans, E., & Ford, F. (2021). Many authors. Journal of Many.")
    files.append((save_doc(doc, "doc4_apa_split.docx"), "a2v"))

    # 5. Vancouver Standard [1]
    doc = Document()
    doc.add_heading("5. Vancouver Standard", level=1)
    doc.add_paragraph("This is vancouver [1].")
    doc.add_paragraph("References")
    doc.add_paragraph("1. Smith A, Jones B. A standard paper. Journal of Testing. 2020;1(1):1-10.")
    files.append((save_doc(doc, "doc5_van_standard.docx"), "v2a"))

    # 6. Vancouver Ranges [1-3]
    doc = Document()
    doc.add_heading("6. Vancouver Ranges", level=1)
    doc.add_paragraph("This is vancouver with a range [1-3].")
    doc.add_paragraph("References")
    doc.add_paragraph("1. Smith A, Jones B. Paper 1. Journal of Testing. 2020;1(1):1-10.")
    doc.add_paragraph("2. Williams C. Paper 2. Journal of Testing. 2021;1(1):1-10.")
    doc.add_paragraph("3. Davis D. Paper 3. Journal of Testing. 2022;1(1):1-10.")
    files.append((save_doc(doc, "doc6_van_ranges.docx"), "v2a"))

    # 7. Vancouver Commas [1, 3, 5]
    doc = Document()
    doc.add_heading("7. Vancouver Commas", level=1)
    doc.add_paragraph("This is vancouver with commas [1, 2].")
    doc.add_paragraph("References")
    doc.add_paragraph("1. Smith A, Jones B. Paper 1. Journal of Testing. 2020;1(1):1-10.")
    doc.add_paragraph("2. Davis D. Paper 3. Journal of Testing. 2022;1(1):1-10.")
    files.append((save_doc(doc, "doc7_van_commas.docx"), "v2a"))

    # 8. Harvard Standard
    doc = Document()
    doc.add_heading("8. Harvard Standard", level=1)
    doc.add_paragraph("This is a standard Harvard citation (Smith and Jones, 2020).")
    doc.add_paragraph("References")
    doc.add_paragraph("Smith, A. and Jones, B. (2020) 'A standard paper', Journal of Testing, 1(1), pp. 1-10.")
    files.append((save_doc(doc, "doc8_harv_standard.docx"), "h2v"))

    # 9. Harvard with specific page numbers
    doc = Document()
    doc.add_heading("9. Harvard Pages", level=1)
    doc.add_paragraph("This has pages (Smith and Jones, 2020, p. 15).")
    doc.add_paragraph("References")
    doc.add_paragraph("Smith, A. and Jones, B. (2020) 'A standard paper', Journal of Testing, 1(1), pp. 1-10.")
    files.append((save_doc(doc, "doc9_harv_pages.docx"), "h2v"))

    # 10. Weird Mixed Content APA
    doc = Document()
    doc.add_heading("10. Mixed Format APA", level=1)
    p = doc.add_paragraph()
    p.add_run("Sometimes people italicize ").bold = True
    p.add_run("(Smith, 2020)").italic = True
    p.add_run(" in the middle of sentences.")
    doc.add_paragraph("References")
    doc.add_paragraph("Smith, A. (2020). A standard paper. Journal of Testing, 1(1), 1-10.")
    files.append((save_doc(doc, "doc10_apa_mixed_formatting.docx"), "a2v"))

    return files

def run_tests():
    files = generate_docs()
    success = 0
    repo_root = Path(__file__).parent.parent
    
    print("Running 10 Edge-Case Conversions...")
    print("-" * 50)
    for i, (fpath, mode) in enumerate(files):
        outpath = fpath.replace(".docx", "_OUTPUT.docx")
        cmd = [
            sys.executable, str(repo_root / "scholarref.py"),
            "--mode", mode,
            "--input", fpath,
            "--output", outpath,
            "--allow-unsupported-parts"
        ]
        
        try:
            res = subprocess.run(cmd, capture_output=True, text=True, cwd=repo_root)
            if res.returncode == 0:
                print(f"[OK] Test {i+1}: {Path(fpath).name} ({mode}) converted successfully.")
                success += 1
            else:
                print(f"[FAIL] Test {i+1}: {Path(fpath).name} failed.")
                print(res.stderr)
        except Exception as e:
            print(f"[ERROR] Test {i+1}: {e}")
            
    print("-" * 50)
    print(f"Passed: {success}/10")
    return success, len(files)


def test_10_wild_cases_cli() -> None:
    success, total = run_tests()
    assert success == total, f"Expected all {total} edge-case CLI conversions to pass, got {success}."

if __name__ == "__main__":
    run_tests()
