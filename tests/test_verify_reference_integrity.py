from __future__ import annotations

import sys
from pathlib import Path

import pytest
from docx import Document

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

import verify_reference_integrity


def _write_minimal_doc(path: Path) -> None:
    doc = Document()
    doc.add_paragraph("Title")
    doc.add_paragraph("References")
    doc.add_paragraph("Smith, J. (2020). Example title. Journal of Tests, 1(1), 1-10.")
    doc.save(path)


def test_full_profile_requires_private_converter_when_missing(monkeypatch, tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"
    _write_minimal_doc(source)
    _write_minimal_doc(output)

    monkeypatch.setattr(verify_reference_integrity, "conv", None)

    with pytest.raises(RuntimeError, match="private local 'convert_to_plosone.py' module"):
        verify_reference_integrity.verify(str(source), str(output), profile="full")


def test_pyproject_does_not_publish_private_plos_entrypoint() -> None:
    pyproject = (ROOT / "pyproject.toml").read_text(encoding="utf-8")
    assert "scholarref-plosone" not in pyproject
    assert "convert_to_plosone" not in pyproject


def test_public_verify_default_profile_is_references_only() -> None:
    assert verify_reference_integrity.verify.__defaults__ == ("references-only",)
