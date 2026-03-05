#!/usr/bin/env python3
"""
Reference integrity verifier for author-date -> Vancouver conversion.

Checks performed:
1) Output reference list numbering is contiguous (1..N).
2) In-text citation numbers are valid and reference-linked.
3) First appearance order of citation numbers is ascending.
4) No APA-style author-year citations remain in body text.
5) Each numbered output reference matches the expected source reference
   on year, lead author identity, and title similarity.
"""

from __future__ import annotations

import argparse
import difflib
import re
import sys
import unicodedata
from dataclasses import dataclass
from typing import Dict, List, Sequence, Tuple

from docx import Document

import scholarref as sref

try:
    import convert_to_plosone as conv
except ModuleNotFoundError:
    conv = None

FIG_CAP_HEAD_RE = re.compile(r"^Fig\.?\s*(\d+)\b", re.IGNORECASE)
FIG_REF_RE = re.compile(r"\bFigs?\.?\s*([0-9][0-9,\sand\-]*)", re.IGNORECASE)
WORD_RE = re.compile(r"[A-Za-z0-9]+(?:[-'][A-Za-z0-9]+)?")


@dataclass
class OutputReference:
    num: int
    line: str
    authors_part: str
    title_part: str
    year: str


def _norm_text(text: str) -> str:
    text = unicodedata.normalize("NFKD", text or "")
    text = text.encode("ascii", "ignore").decode("ascii")
    text = text.lower()
    text = re.sub(r"[^a-z0-9]+", " ", text).strip()
    return re.sub(r"\s+", " ", text)


def _norm_name(text: str) -> str:
    text = _norm_text(text)
    toks = text.split()
    # Drop trailing initials from Vancouver-style lead author chunks.
    while toks and len(toks[-1]) <= 2 and toks[-1].isalpha():
        toks.pop()
    return " ".join(toks)


def _split_authors_and_title(vancouver_rest: str) -> Tuple[str, str]:
    first_dot = vancouver_rest.find(". ")
    if first_dot == -1:
        return vancouver_rest.strip(), ""
    authors = vancouver_rest[:first_dot].strip()
    remainder = vancouver_rest[first_dot + 2 :].strip()
    second_dot = remainder.find(". ")
    if second_dot == -1:
        return authors, remainder.rstrip(".")
    title = remainder[:second_dot].strip().rstrip(".")
    return authors, title


def _find_reference_header(doc: Document) -> int:
    return sref.ensure_reference_header(doc)


def _require_private_full_profile_support():
    if conv is None:
        raise RuntimeError(
            "The 'full' verification profile requires the private local "
            "'convert_to_plosone.py' module, which is not distributed with ScholarRef. "
            "Use '--profile references-only' for the public GitHub build."
        )
    return conv


def _extract_output_references(doc: Document) -> Tuple[int, List[OutputReference], List[str]]:
    issues: List[str] = []
    ref_idx = _find_reference_header(doc)
    if ref_idx == -1:
        return -1, [], ["Could not find reference header in output document."]

    refs: List[OutputReference] = []
    for p in doc.paragraphs[ref_idx + 1 :]:
        line = p.text.strip()
        if not line:
            continue
        m = re.match(r"^(\d+)\.\s+(.*)$", line)
        if not m:
            issues.append(f"Non-numbered reference line: {line[:120]}")
            continue
        num = int(m.group(1))
        rest = m.group(2).strip()
        authors_part, title_part = _split_authors_and_title(rest)
        year = ""
        ym = re.search(r"\.\s*((?:19|20)\d{2}[a-z]?)\s*;", rest)
        if ym:
            year = ym.group(1)
        else:
            ym2 = re.search(r";\s*((?:19|20)\d{2}[a-z]?)\s*\.", rest)
            if ym2:
                year = ym2.group(1)
            else:
                years = re.findall(r"\b(?:19|20)\d{2}[a-z]?\b", rest)
                year = years[-1] if years else ""
        refs.append(
            OutputReference(
                num=num,
                line=line,
                authors_part=authors_part,
                title_part=title_part,
                year=year,
            )
        )
    return ref_idx, refs, issues


def _expand_bracket_token(token: str) -> List[int]:
    token = token.replace("â€“", "-").replace("–", "-")
    out: List[int] = []
    for part in token.split(","):
        part = part.strip()
        if not part:
            continue
        if "-" in part:
            a, b = part.split("-", 1)
            if a.isdigit() and b.isdigit():
                out.extend(range(int(a), int(b) + 1))
        elif part.isdigit():
            out.append(int(part))
    return out


def _extract_citation_numbers(
    body_paragraphs: Sequence,
) -> Tuple[List[int], Dict[int, Tuple[int, int]]]:
    numbers: List[int] = []
    first_pos: Dict[int, Tuple[int, int]] = {}
    for pi, p in enumerate(body_paragraphs):
        t = p.text
        for m in re.finditer(r"\[(\d+(?:[,\-–]\d+)*)\]", t):
            expanded = _expand_bracket_token(m.group(1))
            for n in expanded:
                numbers.append(n)
                if n not in first_pos:
                    first_pos[n] = (pi, m.start())
    return numbers, first_pos


def _first_nonempty_line(text: str) -> str:
    for ln in re.split(r"[\r\n]+", text or ""):
        s = ln.strip()
        if s:
            return s
    return ""


def _extract_figure_numbers(text: str) -> List[int]:
    nums: List[int] = []
    for m in FIG_REF_RE.finditer(text or ""):
        for tok in re.findall(r"\d+", m.group(1)):
            n = int(tok)
            if n not in nums:
                nums.append(n)
    return nums


def _extract_title_from_caption_text(text: str, fig_num: int) -> str:
    lines = [ln.strip() for ln in re.split(r"[\r\n]+", text or "") if ln.strip()]
    if not lines:
        return ""
    first = lines[0]
    m = re.match(rf"^Fig\.?\s*{fig_num}\b(.*)$", first, re.IGNORECASE)
    if not m:
        return ""
    tail = m.group(1).strip(" .:-")
    if tail:
        return tail
    for ln in lines[1:]:
        if ln.lower().startswith("note."):
            break
        if FIG_CAP_HEAD_RE.match(ln):
            break
        if ln:
            return ln.strip(" .")
    return ""


def _figure_checks(paragraphs: Sequence) -> Tuple[List[str], List[str]]:
    failures: List[str] = []
    warnings: List[str] = []

    captions: Dict[int, int] = {}
    cap_idx_set = set()
    for i, p in enumerate(paragraphs):
        head = _first_nonempty_line(p.text)
        m = FIG_CAP_HEAD_RE.match(head)
        if m:
            n = int(m.group(1))
            captions.setdefault(n, i)
            cap_idx_set.add(i)

            if not re.match(rf"^Fig\.?\s*{n}\.\s+", head):
                failures.append(f"Figure caption {n} is not normalized as 'Fig {n}. Title.'")
            title = _extract_title_from_caption_text(p.text, n)
            if not title:
                failures.append(f"Figure caption {n} is missing a title.")
            else:
                wc = len(WORD_RE.findall(title))
                if wc > 15:
                    failures.append(f"Figure caption {n} title exceeds 15 words ({wc}).")
                if not p.text.strip().endswith("."):
                    warnings.append(f"Figure caption {n} does not end with a period.")

    first_cite: Dict[int, int] = {}
    first_cite_groups: Dict[int, List[int]] = {}
    for i, p in enumerate(paragraphs):
        if i in cap_idx_set:
            continue
        para_nums: List[int] = []
        for n in _extract_figure_numbers(p.text):
            if n not in para_nums:
                para_nums.append(n)
        for n in para_nums:
            if n not in first_cite:
                first_cite[n] = i
                first_cite_groups.setdefault(i, []).append(n)

    for n in sorted(captions):
        if n not in first_cite:
            failures.append(f"Fig {n} caption exists but no in-text citation was found.")
            continue
        if first_cite[n] > captions[n]:
            failures.append(
                f"Fig {n} caption appears before first citation (caption paragraph {captions[n]}, "
                f"first citation paragraph {first_cite[n]})."
            )

    for para_idx in sorted(first_cite_groups):
        expected_figs = first_cite_groups[para_idx]
        j = para_idx + 1
        k = 0
        while j < len(paragraphs) and k < len(expected_figs):
            txt = paragraphs[j].text.strip()
            if not txt:
                j += 1
                continue
            head = _first_nonempty_line(txt)
            m = FIG_CAP_HEAD_RE.match(head)
            if m and int(m.group(1)) == expected_figs[k]:
                k += 1
                j += 1
                continue
            if txt.lower().startswith("note."):
                j += 1
                continue
            break
        if k < len(expected_figs):
            for n in expected_figs[k:]:
                failures.append(
                    f"Fig {n} caption is not immediately after first citation block "
                    f"(citation paragraph {para_idx}, caption paragraph {captions.get(n, -1)})."
                )

    order = [n for n, _ in sorted(first_cite.items(), key=lambda kv: kv[1])]
    if order and order != sorted(order):
        failures.append("Figure first-citation order is not ascending.")

    return failures, warnings


def _build_expected_order_from_source(src_doc: Document, profile: str):
    if profile == "references-only":
        ref_idx = _find_reference_header(src_doc)
        if ref_idx == -1:
            raise RuntimeError("Could not find reference header in source document.")
        ref_texts = [p.text.strip() for p in src_doc.paragraphs[ref_idx + 1 :] if p.text.strip()]
        refs = [sref.parse_apa_reference(t) for t in ref_texts]
        sref.assign_year_suffixes(refs)
        engine = sref.ApaToVancouverEngine(refs)
        engine.scan(
            sref.collect_body_paragraphs_before_reference(src_doc, ref_idx)
            + sref.collect_header_footer_paragraphs(src_doc)
        )
        ordered = [engine.by_key[k] for k in engine.order if k in engine.by_key]
        # Mirror ScholarRef default behavior (append uncited references).
        for r in refs:
            if r["key"] not in engine.key2num:
                ordered.append(r)
        return refs, ordered

    # Mirror full manuscript converter preprocessing.
    full_profile = _require_private_full_profile_support()
    full_profile.fix_title(src_doc)
    full_profile.add_short_title(src_doc)
    full_profile.clean_title_page(src_doc)
    full_profile.fix_section_headers(src_doc)
    full_profile.dismantle_declarations(src_doc)

    ref_idx = _find_reference_header(src_doc)
    if ref_idx == -1:
        raise RuntimeError("Could not find reference header in source document.")

    ref_texts = [p.text.strip() for p in src_doc.paragraphs[ref_idx + 1 :] if p.text.strip()]
    refs = [full_profile.parse_ref(t) for t in ref_texts]
    engine = full_profile.CitationEngine(refs)
    engine.scan(src_doc.paragraphs[:ref_idx])
    ordered = [engine.by_key[k] for k in engine.order if k in engine.by_key]
    return refs, ordered


def _title_similarity(a: str, b: str) -> float:
    return difflib.SequenceMatcher(a=_norm_text(a), b=_norm_text(b)).ratio()


def _lead_author_ok(source_ref: dict, out_ref: OutputReference) -> bool:
    src_authors = source_ref.get("authors") or []
    if not src_authors:
        return True
    src_lead = _norm_name(src_authors[0][0])
    out_lead = _norm_name(out_ref.authors_part.split(",")[0])
    if not src_lead or not out_lead:
        return True
    return src_lead in out_lead or out_lead in src_lead


def verify(source_path: str, output_path: str, profile: str = "references-only") -> int:
    failures: List[str] = []
    warnings: List[str] = []

    src_doc = Document(source_path)
    out_doc = Document(output_path)

    _, expected_order = _build_expected_order_from_source(src_doc, profile)

    out_ref_idx, out_refs, out_ref_issues = _extract_output_references(out_doc)
    failures.extend(out_ref_issues)
    if out_ref_idx == -1:
        for f in failures:
            print(f"[FAIL] {f}")
        return 1

    body_paragraphs = (
        sref.collect_body_paragraphs_before_reference(out_doc, out_ref_idx)
        + sref.collect_header_footer_paragraphs(out_doc)
    )
    used_nums, first_pos = _extract_citation_numbers(body_paragraphs)
    used_set = set(used_nums)

    if profile == "full":
        full_profile = _require_private_full_profile_support()
        # Structural section checks around declarations/acknowledgments.
        ack_indices = [i for i, p in enumerate(out_doc.paragraphs) if p.text.strip() == "Acknowledgments"]
        if len(ack_indices) != 1:
            failures.append(f"Expected exactly one 'Acknowledgments' header, found {len(ack_indices)}.")
        decl_markers = [
            "Declarations",
            "Funding",
            "Conflicts of interest",
            "Ethics approval",
            "Consent to participate",
            "Consent for publication",
            "Availability of data and material",
            "Code availability",
            "Authors' contributions",
            "Open Practices Statement",
        ]
        remnant_hits: List[Tuple[int, str]] = []
        for i, p in enumerate(body_paragraphs):
            t = p.text.strip()
            if t and any(t.startswith(m) for m in decl_markers):
                remnant_hits.append((i, t))
        if remnant_hits:
            failures.append(
                "Declarations remnants found in output body: "
                + ", ".join(f"{i}:{txt[:40]}" for i, txt in remnant_hits[:10])
            )

        # Figure placement/caption checks for PLOS ONE conformance.
        out_body_only = out_doc.paragraphs[:out_ref_idx]
        fig_fail, fig_warn = _figure_checks(out_body_only)
        failures.extend(fig_fail)
        warnings.extend(fig_warn)

        # Revised-submission title guidance: 15 words or fewer.
        title_text = full_profile._extract_main_title(out_doc)
        if title_text:
            title_words = len(WORD_RE.findall(title_text))
            if title_words > 15:
                warnings.append(f"Title exceeds 15 words ({title_words}).")

    # Check output reference numbering structure.
    out_nums = [r.num for r in out_refs]
    if out_nums != list(range(1, len(out_nums) + 1)):
        failures.append("Output references are not numbered contiguously from 1..N.")

    # Expected count from source citation order should match output list size.
    if len(out_refs) != len(expected_order):
        failures.append(
            f"Reference count mismatch: output has {len(out_refs)} vs expected {len(expected_order)}."
        )

    # In-text citation linkage checks.
    out_num_set = set(out_nums)
    missing_in_refs = sorted(used_set - out_num_set)
    if missing_in_refs:
        failures.append(f"Citations point to missing references: {missing_in_refs}")

    unused_refs = sorted(out_num_set - used_set)
    if unused_refs:
        failures.append(f"Uncited references remain in output list: {unused_refs}")

    # First-appearance order should be numeric ascending for Vancouver numbering.
    first_order = [n for n, _ in sorted(first_pos.items(), key=lambda kv: kv[1])]
    if first_order != sorted(first_order):
        failures.append("Citation first-appearance order is not ascending by reference number.")

    # No remaining APA-style author-year citations in body.
    apa_pat = re.compile(r"\([A-Z][^)]*,\s*(?:19|20)\d{2}[a-z]?[^)]*\)")
    apa_hits = [i for i, p in enumerate(body_paragraphs) if apa_pat.search(p.text)]
    if apa_hits:
        failures.append(f"Remaining APA-style body citations detected in paragraphs: {apa_hits[:20]}")

    # Catch remaining narrative author-year citations, including page-qualified forms.
    narr_pat = re.compile(
        r"\b[A-Z][A-Za-z'`\-]+(?:\s+(?:&|and)\s+[A-Z][A-Za-z'`\-]+)?"
        r"(?:\s+et\s+al\.)?\s*\((?:19|20)\d{2}[a-z]?(?:,\s*[^)]*)?\)"
    )
    narr_hits = [i for i, p in enumerate(body_paragraphs) if narr_pat.search(p.text)]
    if narr_hits:
        failures.append(f"Remaining narrative author-year citations detected in paragraphs: {narr_hits[:20]}")

    # Vancouver style conformance checks for ScholarRef reference-only profile.
    if profile == "references-only":
        full_name_hits: List[Tuple[int, str]] = []
        full_names = [
            "The Journal of Positive Psychology",
            "Journal of Toxicology and Environmental Health, Part B",
            "Psychological Science",
            "Human Factors: The Journal of the Human Factors and Ergonomics Society",
            "Behavioral Sciences",
            "Experimental Psychology",
            "International Journal of Psychophysiology",
        ]
        for out_ref in out_refs:
            line = out_ref.line
            if "et al.." in line:
                failures.append(f"Double period after et al. at ref {out_ref.num}.")
            if re.search(r"\.\s*[‘'\"“].+[’'\"”](?:,|\.)\s", line):
                failures.append(f"Quoted title wrapper remains at ref {out_ref.num}.")
            if re.search(r",\s*\d+\([^)]*\),\s*(?:pp?\.)?\s*[^;]*;\s*(?:19|20)\d{2}", line):
                failures.append(f"Harvard-style reference structure remains at ref {out_ref.num}.")
            if re.search(r":\s*pp?\.\s*(?:\.|$)", line):
                failures.append(f"Empty p./pp. page placeholder at ref {out_ref.num}.")
            if "?. " in line or "!. " in line:
                failures.append(f"Question/exclamation punctuation artifact at ref {out_ref.num}.")
            if re.search(r":\d+\?\d+\b", line):
                failures.append(f"Page-range separator artifact ('?') at ref {out_ref.num}.")
            if "a journal of the Association for." in line:
                failures.append(f"Journal field appears split/truncated at ref {out_ref.num}.")
            if any(name in line for name in full_names):
                full_name_hits.append((out_ref.num, line))

        if full_name_hits:
            failures.append(
                "Unabbreviated journal names remain at refs: "
                + ", ".join(str(n) for n, _ in full_name_hits[:15])
            )

    # Per-reference source integrity checks.
    for out_ref in out_refs:
        if not (1 <= out_ref.num <= len(expected_order)):
            failures.append(f"Reference number {out_ref.num} is outside expected range.")
            continue
        src_ref = expected_order[out_ref.num - 1]
        src_year = f"{src_ref.get('year', '')}{src_ref.get('ysuf', '')}"
        if src_year and out_ref.year and src_year != out_ref.year:
            failures.append(
                f"Year mismatch at ref {out_ref.num}: source={src_year} output={out_ref.year}"
            )

        if not _lead_author_ok(src_ref, out_ref):
            src_lead = (src_ref.get("authors") or [("", "")])[0][0]
            failures.append(
                f"Lead author mismatch at ref {out_ref.num}: source='{src_lead}' output='{out_ref.authors_part}'"
            )

        src_title = (src_ref.get("title", "") or "").strip()
        out_title = (out_ref.title_part or "").strip()
        if src_title and out_title:
            sim = _title_similarity(src_title, out_title)
            # For chapter-style source titles ("X. In ..."), allow match on the leading chapter title.
            if sim < 0.70 and ". In " in src_title:
                lead_title = src_title.split(". In ", 1)[0].strip()
                if lead_title:
                    sim = max(sim, _title_similarity(lead_title, out_title))
            if sim < 0.55:
                failures.append(
                    f"Low title similarity at ref {out_ref.num} ({sim:.2f}): "
                    f"source='{src_title[:90]}' output='{out_title[:90]}'"
                )
            elif sim < 0.70:
                warnings.append(
                    f"Moderate title similarity at ref {out_ref.num} ({sim:.2f})"
                )
        else:
            warnings.append(
                f"Missing title field for similarity check at ref {out_ref.num}"
            )

    print("Reference Integrity Verification")
    print("- source:", source_path)
    print("- output:", output_path)
    print("- profile:", profile)
    print(f"- output references: {len(out_refs)}")
    print(f"- in-text citation numbers used: {len(used_set)}")

    if warnings:
        print(f"- warnings: {len(warnings)}")
        for w in warnings[:15]:
            print("  [WARN]", w)
        if len(warnings) > 15:
            print(f"  ... {len(warnings) - 15} more warnings")

    if failures:
        print(f"- failures: {len(failures)}")
        for f in failures:
            print("  [FAIL]", f)
        return 1

    print("- result: PASS (no integrity failures)")
    return 0


def main() -> int:
    parser = argparse.ArgumentParser(description="Verify converted reference integrity.")
    parser.add_argument("--source", default="manuscript submission 1.docx")
    parser.add_argument("--output", default="manuscript_PLOSONE.docx")
    parser.add_argument(
        "--profile",
        choices=["full", "references-only"],
        default="references-only",
        help=(
            "Use 'references-only' for public ScholarRef outputs. "
            "'full' requires the private local convert_to_plosone.py module."
        ),
    )
    args = parser.parse_args()
    return verify(args.source, args.output, profile=args.profile)


if __name__ == "__main__":
    raise SystemExit(main())

