#!/usr/bin/env python3
"""
Bidirectional reference-style converter for .docx manuscripts.

Modes:
- APA 7 author-date -> Vancouver numeric
- Vancouver numeric -> APA 7 author-date (best effort)
"""

from __future__ import annotations

import argparse
import re
import sys
import unicodedata
from pathlib import Path
from typing import Dict, List, Optional, Sequence, Tuple

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")


JABBR: Dict[str, str] = {
    "Annual Review of Neuroscience": "Annu Rev Neurosci",
    "Journal of Environmental Psychology": "J Environ Psychol",
    "Sports Medicine": "Sports Med",
    "Proceedings of the Human Factors and Ergonomics Society Annual Meeting":
        "Proc Hum Factors Ergon Soc Annu Meet",
    "Brain, Behavior, & Immunity - Health": "Brain Behav Immun Health",
    "Brain, Behavior, \u0026 Immunity - Health": "Brain Behav Immun Health",
    "Proceedings of the National Academy of Sciences of the United States of America":
        "Proc Natl Acad Sci U S A",
    "Proceedings of the National Academy of Sciences": "Proc Natl Acad Sci U S A",
    "Biological Psychology": "Biol Psychol",
    "Cortex": "Cortex",
    "Journal of Cognitive Neuroscience": "J Cogn Neurosci",
    "Memory & Cognition": "Mem Cognit",
    "Memory \u0026 Cognition": "Mem Cognit",
    "Disability and Rehabilitation": "Disabil Rehabil",
    "Frontiers in Psychology": "Front Psychol",
    "Current Directions in Psychological Science": "Curr Dir Psychol Sci",
    "Medicine & Science in Sports & Exercise": "Med Sci Sports Exerc",
    "Medicine \u0026 Science in Sports \u0026 Exercise": "Med Sci Sports Exerc",
    "PLOS ONE": "PLoS One",
    "PLoS ONE": "PLoS One",
    "Behavior Research Methods": "Behav Res Methods",
    "Behavior Research Methods, Instruments, & Computers":
        "Behav Res Methods Instrum Comput",
    "Behavior Research Methods, Instruments, \u0026 Computers":
        "Behav Res Methods Instrum Comput",
    "Journal of Experimental Psychology: Learning, Memory, and Cognition":
        "J Exp Psychol Learn Mem Cogn",
    "Journal of Experimental Psychology: General": "J Exp Psychol Gen",
    "Brain": "Brain",
    "Brain and Cognition": "Brain Cogn",
    "Journal of Statistical Software": "J Stat Softw",
    "Current Opinion in Neurobiology": "Curr Opin Neurobiol",
    "Psychological Bulletin": "Psychol Bull",
    "Attention, Perception, & Psychophysics": "Atten Percept Psychophys",
    "Attention, Perception, \u0026 Psychophysics": "Atten Percept Psychophys",
    "Acta Psychologica": "Acta Psychol (Amst)",
    "Psychophysiology": "Psychophysiology",
    "Nature Communications": "Nat Commun",
    "Nature": "Nature",
    "Current Biology": "Curr Biol",
    "Trends in Cognitive Sciences": "Trends Cogn Sci",
    "Frontiers in Human Neuroscience": "Front Hum Neurosci",
    "Frontiers in Neurology": "Front Neurol",
    "Human Brain Mapping": "Hum Brain Mapp",
    "Cognitive, Affective, & Behavioral Neuroscience": "Cogn Affect Behav Neurosci",
    "Cognitive, Affective, \u0026 Behavioral Neuroscience": "Cogn Affect Behav Neurosci",
    "Perspectives on Psychological Science": "Perspect Psychol Sci",
    "Perspectives on psychological science : a journal of the Association for Psychological Science":
        "Perspect Psychol Sci",
    "Perspectives on psychological science: a journal of the Association for Psychological Science":
        "Perspect Psychol Sci",
    "Brain Sciences": "Brain Sci",
    "Neuropsychologia": "Neuropsychologia",
    "Cognition": "Cognition",
    "Scientific Reports": "Sci Rep",
    "Psychometrika": "Psychometrika",
    "British Journal of Psychology": "Br J Psychol",
    "Journal of Memory and Language": "J Mem Lang",
    "The Quarterly Journal of Experimental Psychology Section A": "Q J Exp Psychol A",
    "The Quarterly Journal of Experimental Psychology": "Q J Exp Psychol",
    "Neural Plasticity": "Neural Plast",
    "Brain and Behavior": "Brain Behav",
    "Psychology of Sport and Exercise": "Psychol Sport Exerc",
    "Canadian Journal of Experimental Psychology / Revue Canadienne de Psychologie Expérimentale":
        "Can J Exp Psychol",
    "Canadian Journal of Experimental Psychology": "Can J Exp Psychol",
    "Behavioral and Brain Functions": "Behav Brain Funct",
    "The Journal of Neuroscience": "J Neurosci",
    "Journal of Neuroscience": "J Neurosci",
    "International Journal of Sports Medicine": "Int J Sports Med",
    "Psychological Review": "Psychol Rev",
    "Memory": "Memory",
    "Motivation Science": "Motiv Sci",
    "Methods in Ecology and Evolution": "Methods Ecol Evol",
    "Journal of Open Source Software": "J Open Source Softw",
    "The Journal of Positive Psychology": "J Posit Psychol",
    "Journal of Toxicology and Environmental Health, Part B": "J Toxicol Environ Health B",
    "Psychological Science": "Psychol Sci",
    "Human Factors: The Journal of the Human Factors and Ergonomics Society": "Hum Factors",
    "Forests": "Forests (Basel)",
    "Behavioral Sciences": "Behav Sci (Basel)",
    "Experimental Psychology": "Exp Psychol",
    "iScience": "iScience",
    "International Journal of Psychophysiology": "Int J Psychophysiol",
    "Journal of Cognitive Psychology": "J Cogn Psychol",
}


def full_text(para) -> str:
    return "".join(r.text for r in para.runs)


def replace_in_runs(para, replacements: Sequence[Tuple[int, int, str]]) -> None:
    if not replacements or not para.runs:
        return
    for rs, re_, nt in sorted(replacements, key=lambda x: x[0], reverse=True):
        runs = list(para.runs)
        bounds: List[Tuple[int, int, int]] = []
        pos = 0
        for i, r in enumerate(runs):
            L = len(r.text)
            bounds.append((pos, pos + L, i))
            pos += L
        affected = [(s, e, idx) for s, e, idx in bounds if s < re_ and e > rs]
        if not affected:
            continue
        affected.sort()
        fs, _fe, fi = affected[0]
        rel_s = rs - fs
        if len(affected) == 1:
            rel_e = re_ - fs
            runs[fi].text = runs[fi].text[:rel_s] + nt + runs[fi].text[rel_e:]
        else:
            runs[fi].text = runs[fi].text[:rel_s] + nt
            for _, _, mi in affected[1:-1]:
                runs[mi].text = ""
            ls, _le, li = affected[-1]
            runs[li].text = runs[li].text[re_ - ls:]


def _norm_space(text: str) -> str:
    return re.sub(r"\s+", " ", (text or "")).strip()


def _norm_text(text: str) -> str:
    t = unicodedata.normalize("NFKD", text or "")
    t = t.encode("ascii", "ignore").decode("ascii")
    t = t.lower()
    t = re.sub(r"[^a-z0-9]+", " ", t).strip()
    return re.sub(r"\s+", " ", t)


REFERENCE_HEADER_LABELS = {
    "references",
    "reference list",
    "bibliography",
    "works cited",
    "literature cited",
}


def normalized_reference_header_text(text: str) -> str:
    return re.sub(r"\s+", " ", re.sub(r"[^A-Za-z0-9]+", " ", text or "").strip()).strip().lower()


def is_reference_header_text(text: str) -> bool:
    return normalized_reference_header_text(text) in REFERENCE_HEADER_LABELS


def _find_ref_headers(doc: Document) -> List[int]:
    out: List[int] = []
    for i, p in enumerate(doc.paragraphs):
        if is_reference_header_text(p.text):
            out.append(i)
    return out


def _find_ref_header(doc: Document, occurrence: int = 1) -> int:
    headers = _find_ref_headers(doc)
    if occurrence < 1 or occurrence > len(headers):
        return -1
    return headers[occurrence - 1]


def _reference_entry_score(text: str) -> int:
    s = _norm_space(text)
    if not s:
        return 0

    score = 0
    if re.match(r"^\[?\d+\]?[\.\)]\s+", s):
        score += 4
    if re.match(r"^[A-Z][A-Za-z'`\-]+,\s*[A-Z]", s):
        score += 2
    if re.match(r"^[A-Z][A-Za-z'`\-]+(?:\s+[A-Z][A-Za-z'`\-]+)?\s+[A-Z]{1,4}[.,]", s):
        score += 2
    if re.search(r"\((?:19|20)\d{2}[a-z]?(?:,\s*[^)]*)?\)", s):
        score += 2
    if re.search(r"(?<!\d)(?:19|20)\d{2}[a-z]?\s*;\s*\d+", s):
        score += 3
    if re.search(r"\bpp?\.\s*\d", s, flags=re.I):
        score += 2
    if re.search(r"\d+\([^)]*\)\s*[:,]\s*\d", s):
        score += 2
    if re.search(r"https?://|doi:", s, flags=re.I):
        score += 1
    if re.search(r"\bet al\.?\b", s, flags=re.I):
        score += 1
    if len(s) >= 40:
        score += 1
    return score


def _citation_key_parts(key: str) -> Tuple[str, str]:
    if ", " not in key:
        return key, ""
    a, y = key.rsplit(", ", 1)
    return a, y


def _citation_base_key(key: str) -> str:
    a, y = _citation_key_parts(key)
    yb = re.sub(r"[a-z]$", "", y)
    return f"{a}, {yb}"


def _assign_year_suffixes(refs: Sequence[dict]) -> None:
    groups: Dict[str, List[dict]] = {}
    for r in refs:
        yr = r.get("year", "")
        if not yr:
            continue
        base = _citation_key(r.get("authors", []), yr, "")
        groups.setdefault(base, []).append(r)

    for base, items in groups.items():
        if len(items) <= 1:
            # Keep existing suffix if present.
            if items and items[0].get("key"):
                items[0]["key"] = _citation_key(
                    items[0].get("authors", []),
                    items[0].get("year", ""),
                    items[0].get("ysuf", ""),
                )
            continue

        existing = [str(r.get("ysuf", "")).strip() for r in items]
        if all(existing) and len(set(existing)) == len(existing):
            for r in items:
                r["key"] = _citation_key(r.get("authors", []), r.get("year", ""), r.get("ysuf", ""))
            continue

        # Deterministic assignment by normalized title, then raw reference line.
        ordered = sorted(
            items,
            key=lambda r: (
                _norm_text(r.get("title", "")),
                _norm_text(r.get("raw", "")),
            ),
        )
        for i, r in enumerate(ordered):
            r["ysuf"] = chr(ord("a") + i)
            r["key"] = _citation_key(r.get("authors", []), r.get("year", ""), r.get("ysuf", ""))


def assign_year_suffixes(refs: Sequence[dict]) -> None:
    _assign_year_suffixes(refs)


def _reference_identity_signature(ref: dict) -> Tuple[object, ...]:
    authors = tuple(
        (_norm_text(sn), _norm_text(ini))
        for sn, ini in (ref.get("authors", []) or [])
    )
    return (
        authors,
        _norm_text(ref.get("year", "") or ""),
        _norm_text(ref.get("ysuf", "") or ""),
        _norm_text(ref.get("title", "") or ref.get("title_part", "") or ""),
        _norm_text(ref.get("journal", "") or ""),
        _norm_text(ref.get("vol", "") or ""),
        _norm_text(ref.get("issue", "") or ""),
        _norm_text(ref.get("pages", "") or ""),
        _norm_text(ref.get("doi", "") or ""),
        bool(ref.get("is_book")),
        _norm_text(ref.get("raw", "") or ""),
    )


def _collapse_exact_duplicate_references(refs: Sequence[dict]) -> Tuple[List[dict], int]:
    seen: Dict[Tuple[object, ...], dict] = {}
    out: List[dict] = []
    duplicates = 0
    for ref in refs:
        sig = _reference_identity_signature(ref)
        if sig in seen:
            duplicates += 1
            continue
        seen[sig] = ref
        out.append(ref)
    return out, duplicates


def _iter_package_xml_parts(doc: Document):
    for part in doc.part.package.parts:
        name = str(getattr(part, "partname", ""))
        if not name.lower().endswith(".xml"):
            continue
        blob = getattr(part, "blob", None)
        if blob is None:
            continue
        try:
            text = blob.decode("utf-8", errors="ignore")
        except Exception:
            continue
        yield name, text


def preflight_docx(
    doc: Document,
    *,
    ref_header_n: int = 1,
    allow_field_codes: bool = False,
    allow_unsupported_parts: bool = False,
) -> Dict[str, object]:
    warnings: List[str] = []
    failures: List[str] = []

    headers = _find_ref_headers(doc)
    auto_ref_start = -1
    if not headers:
        auto_ref_start = auto_detect_reference_start(doc)
        if auto_ref_start == -1:
            failures.append(
                "Could not find reference header or infer a reference list automatically."
            )
        else:
            headers = [auto_ref_start]
            warnings.append(
                f"No explicit reference header found; auto-detected reference list starting at paragraph index {auto_ref_start}."
            )
    elif ref_header_n < 1 or ref_header_n > len(headers):
        failures.append(
            f"Requested reference header #{ref_header_n} but only {len(headers)} reference header(s) found."
        )
    elif len(headers) > 1:
        warnings.append(
            f"Multiple reference headers found at paragraph indices {headers}; using header #{ref_header_n}."
        )

    addin_hits: List[str] = []
    textbox_cite_parts: List[str] = []
    foot_endnote_cite_parts: List[str] = []
    cite_like = re.compile(
        r"\[(\d+(?:[,\-–]\d+)*)\]|\([A-Z][^)]*,\s*(?:19|20)\d{2}[a-z]?[^)]*\)"
    )
    for name, xml in _iter_package_xml_parts(doc):
        if re.search(
            r"ADDIN\s+(?:ZOTERO|EN\.CITE|MENDELEY|PAPERPILE)|CSL_CITATION|CSL_BIBLIOGRAPHY",
            xml,
            flags=re.I,
        ):
            addin_hits.append(name)

        if "<w:txbxContent" in xml and cite_like.search(xml):
            textbox_cite_parts.append(name)

        lname = name.lower()
        if ("footnotes.xml" in lname or "endnotes.xml" in lname) and cite_like.search(xml):
            foot_endnote_cite_parts.append(name)

    if addin_hits and not allow_field_codes:
        failures.append(
            "Reference-manager field codes detected (e.g., Zotero/EndNote/Mendeley). "
            "Please unlink fields or export plain-text citations first. Parts: "
            + ", ".join(sorted(set(addin_hits))[:8])
        )
    elif addin_hits:
        warnings.append("Field-code citations detected and allowed by option override.")

    if textbox_cite_parts:
        msg = (
            "Citation-like content found in text boxes; those regions are not converted by the current engine. Parts: "
            + ", ".join(sorted(set(textbox_cite_parts))[:8])
        )
        if allow_unsupported_parts:
            warnings.append(msg)
        else:
            failures.append(msg)

    if foot_endnote_cite_parts:
        msg = (
            "Citation-like content found in footnotes/endnotes; those regions are not converted by the current engine. Parts: "
            + ", ".join(sorted(set(foot_endnote_cite_parts))[:8])
        )
        if allow_unsupported_parts:
            warnings.append(msg)
        else:
            failures.append(msg)

    return {
        "warnings": warnings,
        "failures": failures,
        "ref_headers": headers,
        "auto_ref_start": auto_ref_start,
    }
    return -1


def _set_references_header(doc: Document, ref_idx: int) -> None:
    if ref_idx < 0 or ref_idx >= len(doc.paragraphs):
        return
    p = doc.paragraphs[ref_idx]
    txt = full_text(p)
    if is_reference_header_text(txt) and txt.strip() != "References":
        p.text = "References"


def _para_style_name(para) -> str:
    try:
        return (getattr(getattr(para, "style", None), "name", "") or "").strip().lower()
    except Exception:
        return ""


def _looks_like_reference_entry(text: str) -> bool:
    return _reference_entry_score(text) >= 3


def _looks_like_reference_section_break(para) -> bool:
    text = _norm_space(full_text(para))
    if not text:
        return False

    style_name = _para_style_name(para)
    if any(token in style_name for token in ("heading", "title", "subtitle")):
        return True

    if re.match(
        r"^(?:appendix|appendices|supplement(?:ary|al)?(?: materials?)?|supporting information|"
        r"online resources?|figure legends?|acknowledg?ments?|funding|declarations?|"
        r"author contributions?|conflicts? of interest|ethics(?: statement)?|"
        r"data availability|supplementary checklist|checklist)\b",
        text,
        flags=re.I,
    ):
        return True

    words = re.findall(r"[A-Za-z0-9][A-Za-z0-9'&/\-]*", text)
    if (
        words
        and len(words) <= 8
        and len(text) <= 80
        and not re.search(r"[.:;!?]|\((?:19|20)\d{2}|https?://|doi:", text, flags=re.I)
    ):
        capitalized = sum(1 for w in words if w[:1].isupper())
        if capitalized >= max(1, len(words) - 1):
            return True
    return False


def _next_nonempty_paragraph(doc: Document, start_idx: int):
    for i in range(start_idx + 1, len(doc.paragraphs)):
        para = doc.paragraphs[i]
        if full_text(para).strip():
            return para
    return None


def _collect_ref_paragraphs_from_start(doc: Document, start_idx: int) -> List[Tuple[int, str]]:
    items: List[Tuple[int, str]] = []
    for i in range(start_idx, len(doc.paragraphs)):
        para = doc.paragraphs[i]
        t = full_text(para).strip()
        if not t:
            continue

        if items and _looks_like_reference_section_break(para):
            break

        if not _looks_like_reference_entry(t):
            next_para = _next_nonempty_paragraph(doc, i)
            next_text = full_text(next_para).strip() if next_para is not None else ""
            if items and (
                next_para is None
                or _looks_like_reference_section_break(next_para)
                or not _looks_like_reference_entry(next_text)
            ):
                break
            if not items:
                continue

        items.append((i, t))
    return items


def _collect_ref_paragraphs(doc: Document, ref_idx: int) -> List[Tuple[int, str]]:
    return _collect_ref_paragraphs_from_start(doc, ref_idx + 1)


def auto_detect_reference_start(doc: Document) -> int:
    candidates: List[Tuple[int, int, int]] = []
    total_paras = len(doc.paragraphs)
    if total_paras == 0:
        return -1

    for i, para in enumerate(doc.paragraphs):
        text = full_text(para).strip()
        if _reference_entry_score(text) < 3:
            continue

        block = _collect_ref_paragraphs_from_start(doc, i)
        if not block:
            continue

        scores = [_reference_entry_score(item_text) for _, item_text in block]
        strong = sum(1 for score in scores if score >= 4)
        if len(block) >= 2 and strong >= 2:
            candidates.append((len(block), sum(scores), i))
            continue

        if (
            len(block) == 1
            and scores[0] >= 5
            and i >= max(0, total_paras - 4)
        ):
            candidates.append((len(block), sum(scores), i))

    if not candidates:
        return -1

    candidates.sort(key=lambda item: (item[0], item[1], item[2]), reverse=True)
    return candidates[0][2]


def _insert_plain_paragraph_before(doc: Document, idx: int, text: str) -> None:
    if not (0 <= idx < len(doc.paragraphs)):
        doc.add_paragraph(text)
        return

    anchor = doc.paragraphs[idx]._element
    new_p = OxmlElement("w:p")
    run_el = OxmlElement("w:r")
    t_el = OxmlElement("w:t")
    t_el.text = text
    t_el.set(qn("xml:space"), "preserve")
    run_el.append(t_el)
    new_p.append(run_el)
    anchor.addprevious(new_p)


def ensure_reference_header(doc: Document, occurrence: int = 1) -> int:
    ref_idx = _find_ref_header(doc, occurrence=occurrence)
    if ref_idx != -1:
        return ref_idx

    auto_start = auto_detect_reference_start(doc)
    if auto_start == -1:
        return -1

    _insert_plain_paragraph_before(doc, auto_start, "References")
    return auto_start


def _remove_paragraphs_by_indices(doc: Document, indices: Sequence[int]) -> None:
    for idx in sorted(set(indices), reverse=True):
        if 0 <= idx < len(doc.paragraphs):
            el = doc.paragraphs[idx]._element
            el.getparent().remove(el)


def _insert_reference_lines(doc: Document, ref_idx: int, lines: Sequence[str]) -> None:
    anchor = doc.paragraphs[ref_idx]._element
    for line in reversed(lines):
        new_p = OxmlElement("w:p")
        run_el = OxmlElement("w:r")
        t_el = OxmlElement("w:t")
        t_el.text = line
        t_el.set(qn("xml:space"), "preserve")
        run_el.append(t_el)
        new_p.append(run_el)
        anchor.addnext(new_p)


def _iter_container_paragraphs(container) -> List:
    if hasattr(container, "element") and hasattr(container.element, "body"):
        root = container.element.body
    elif hasattr(container, "_element"):
        root = container._element
    elif hasattr(container, "_tc"):
        root = container._tc
    else:
        return []

    out: List = []
    for child in root.iterchildren():
        if child.tag == qn("w:p"):
            out.append(Paragraph(child, container))
        elif child.tag == qn("w:tbl"):
            tbl = Table(child, container)
            seen_cells = set()
            for row in tbl.rows:
                for cell in row.cells:
                    cell_id = id(cell._tc)
                    if cell_id in seen_cells:
                        continue
                    seen_cells.add(cell_id)
                    out.extend(_iter_container_paragraphs(cell))
    return out


def collect_body_paragraphs_before_reference(doc: Document, ref_idx: Optional[int] = None) -> List:
    if ref_idx is None:
        ref_idx = _find_ref_header(doc)
    ref_el = doc.paragraphs[ref_idx]._element if 0 <= ref_idx < len(doc.paragraphs) else None

    body = doc.element.body
    out: List = []
    for child in body.iterchildren():
        if ref_el is not None and child == ref_el:
            break
        if child.tag == qn("w:p"):
            out.append(Paragraph(child, doc))
        elif child.tag == qn("w:tbl"):
            tbl = Table(child, doc)
            seen_cells = set()
            for row in tbl.rows:
                for cell in row.cells:
                    cell_id = id(cell._tc)
                    if cell_id in seen_cells:
                        continue
                    seen_cells.add(cell_id)
                    out.extend(_iter_container_paragraphs(cell))
    return out


def collect_header_footer_paragraphs(doc: Document) -> List:
    out: List = []
    seen = set()
    for sec in doc.sections:
        parts = [
            sec.header,
            sec.footer,
            sec.first_page_header,
            sec.first_page_footer,
            sec.even_page_header,
            sec.even_page_footer,
        ]
        for part in parts:
            el = getattr(part, "_element", None)
            if el is None:
                continue
            pid = id(el)
            if pid in seen:
                continue
            seen.add(pid)
            out.extend(_iter_container_paragraphs(part))
    return out


def _extract_doi(text: str) -> str:
    dm = re.search(r"\bdoi:\s*([^\s]+)", text, flags=re.I)
    if dm:
        return dm.group(1).rstrip(".")
    um = re.search(r"https?://(?:dx\.)?doi\.org/([^\s]+)", text, flags=re.I)
    if um:
        return um.group(1).rstrip(".")
    return ""


def _strip_surrounding_quotes(text: str) -> str:
    t = text.strip().strip(",").strip()
    pairs = [("‘", "’"), ("“", "”"), ("'", "'"), ('"', '"')]
    for left, right in pairs:
        if t.startswith(left) and t.endswith(right) and len(t) >= 2:
            t = t[1:-1].strip()
            break
    t = t.strip(" '\"“”‘’")
    return t


def _clean_pages(text: str) -> str:
    t = (text or "").strip().rstrip(".;,")
    t = re.sub(r"^\s*pp?\.\s*", "", t, flags=re.I).strip()
    t = t.replace("−", "-").replace("–", "-")
    t = re.sub(r"(?<=\d)\?(?=\d)", "-", t)
    t = re.sub(r"\s*-\s*", "-", t)
    t = re.sub(r"\s+", " ", t).strip()
    return t


def _end_punct(text: str) -> str:
    t = (text or "").strip()
    if not t:
        return ""
    if t[-1] in ".?!":
        return t
    return t + "."


def _split_title_and_source(after_clean: str) -> Tuple[str, str]:
    s = _norm_space(after_clean).strip().rstrip(".")
    if not s:
        return "", ""

    if s[0] in "‘“'\"":
        close_map = {"‘": "’", "“": "”", "'": "'", '"': '"'}
        close_char = close_map.get(s[0], s[0])
        close_positions = [i for i, ch in enumerate(s) if ch == close_char and i > 0]
        for close_idx in reversed(close_positions):
            rest0 = s[close_idx + 1 :].lstrip()
            if rest0.startswith(","):
                title = s[1:close_idx].strip()
                rest = rest0[1:].lstrip(" ,.;")
                return title, rest
            if rest0.startswith("."):
                title = s[1:close_idx].strip()
                rest = rest0[1:].lstrip(" ,.;")
                if rest:
                    return title, rest
        if close_positions:
            close_idx = close_positions[-1]
            title = s[1:close_idx].strip()
            rest = s[close_idx + 1 :].lstrip(" ,.;")
            return title, rest

    for jn in sorted(JABBR, key=len, reverse=True):
        pos = s.find(jn)
        if pos > 0:
            title = s[:pos].rstrip(" ,.;")
            rest = s[pos:].lstrip(" ,.;")
            if title:
                return title, rest

    if ". " in s:
        title, rest = s.rsplit(". ", 1)
        return title.strip(), rest.strip()
    return s.strip(), ""


def _parse_source_fields(source: str) -> Tuple[str, str, str, str, bool]:
    src = _norm_space(source).strip().rstrip(".")
    if not src:
        return "", "", "", "", False

    is_book = bool(re.search(r"\b(?:edn|edition|press|associates|publisher)\b", src, flags=re.I))
    journal = src
    vol = ""
    issue = ""
    pages = ""

    # Capture right-most volume/issue/pages pattern while allowing commas in journal name.
    vm = re.search(r",\s*(\d+)\s*(?:\(([^)]+)\))?(?:,\s*(.*))?$", src)
    if vm:
        journal = src[: vm.start()].strip(" ,.")
        vol = (vm.group(1) or "").strip()
        issue = (vm.group(2) or "").strip()
        pages = _clean_pages(vm.group(3) or "")
    else:
        pm = re.search(r",\s*(pp?\.\s*.+)$", src, flags=re.I)
        if pm:
            journal = src[: pm.start()].strip(" ,.")
            pages = _clean_pages(pm.group(1))
        else:
            journal = src.strip(" ,.")

    if pages in {"p", "pp"}:
        pages = ""

    return journal, vol, issue, pages, is_book


def _parse_apa_authors(s: str) -> List[Tuple[str, str]]:
    s = re.sub(r"\s+", " ", s.strip()).rstrip(".")
    if not s:
        return []

    def _is_initials(token: str) -> bool:
        compact = re.sub(r"[\s\.\-]", "", token.strip())
        return bool(compact) and compact.isalpha() and compact.upper() == compact

    def _norm_initials(token: str) -> str:
        letters = re.findall(r"[A-Z]", token or "")
        return " ".join(f"{ch}." for ch in letters)

    # Handle compact Harvard-style "Surname, X.Y. et al." author strings.
    etal_match = re.match(r"^(.*?)(?:,?\s+)?et al\.?$", s, flags=re.I)
    if etal_match:
        lead = etal_match.group(1).strip().rstrip(",")
        if "," in lead:
            sn, ini = lead.split(",", 1)
            sn = sn.strip()
            ini_norm = _norm_initials(ini)
            if sn:
                return [(sn, ini_norm), ("et al.", "")]
        if lead:
            return [(lead, ""), ("et al.", "")]

    normalized = re.sub(r"\s+(?:&|and)\s+", ", ", s)
    parts = [p.strip() for p in normalized.split(",") if p.strip()]
    authors: List[Tuple[str, str]] = []
    i = 0
    while i < len(parts):
        surname = parts[i]
        if i + 1 < len(parts) and _is_initials(parts[i + 1]):
            authors.append((surname, _norm_initials(parts[i + 1])))
            i += 2
        else:
            if len(parts) == 1:
                return [(parts[0].replace("&", "").strip(), "")]
            authors.append((surname, ""))
            i += 1
    if all(not ini for _, ini in authors):
        return [(s.replace("&", "").strip(), "")]
    return authors


def _citation_key(authors: Sequence[Tuple[str, str]], year: str, suffix: str) -> str:
    if not authors:
        return f"Unknown, {year}{suffix}"
    sn = [a[0] for a in authors]

    if len(sn) == 1:
        if re.search(r"\bet al\b", sn[0], flags=re.I):
            lead = re.sub(r"\bet al\.?", "", sn[0], flags=re.I).strip(" ,")
            author_part = f"{lead} et al."
        else:
            author_part = sn[0]
    elif len(sn) == 2:
        if any(re.search(r"\bet al\b", x, flags=re.I) for x in sn):
            lead = re.sub(r"\bet al\.?", "", sn[0], flags=re.I).strip(" ,")
            author_part = f"{lead} et al."
        else:
            author_part = f"{sn[0]} & {sn[1]}"
    else:
        author_part = f"{sn[0]} et al."
    return f"{author_part}, {year}{suffix}"


def parse_author_date_reference(text: str) -> dict:
    text = text.replace("\u00A0", " ")
    text = re.sub(r"\s+", " ", text).strip()
    doi = _extract_doi(text)
    text = re.sub(r"\s*https?://\S+\s*$", "", text).strip()
    text = re.sub(r"\s*doi:\s*[^\s]+\s*$", "", text, flags=re.I).strip()
    ref = dict(
        raw=text,
        authors=[],
        year="",
        ysuf="",
        title="",
        journal="",
        vol="",
        issue="",
        pages="",
        key="",
        num=None,
        is_book=False,
        is_nature=False,
        doi=doi,
    )

    nm = re.search(
        r"^(.+?)\.\s*(Nature)\s+(\d+)\s*,\s*([^\s()]+)\s*\((\d{4})\)\.?$",
        text,
    )
    if nm:
        pre = nm.group(1).strip()
        if ". " in pre:
            auth_raw, title = pre.rsplit(". ", 1)
        else:
            auth_raw, title = pre, ""
        ref["authors"] = _parse_apa_authors(auth_raw)
        ref["title"] = title.strip().rstrip(".")
        ref["journal"] = nm.group(2)
        ref["vol"] = nm.group(3)
        ref["pages"] = nm.group(4)
        ref["year"] = nm.group(5)
        ref["is_nature"] = True
        ref["key"] = _citation_key(ref["authors"], ref["year"], "")
        return ref

    ym = re.search(r"\((\d{4})([a-z]?)(?:,\s*[^)]*)?\)", text)
    if not ym:
        ym2 = re.search(r"\b((?:19|20)\d{2})([a-z])?\b", text)
        if ym2:
            ref["year"] = ym2.group(1)
            ref["ysuf"] = ym2.group(2) or ""
        ref["key"] = text[:60]
        return ref

    ref["year"] = ym.group(1)
    ref["ysuf"] = ym.group(2) or ""
    auth_raw = text[: ym.start()].strip().rstrip(".")
    ref["authors"] = _parse_apa_authors(auth_raw)
    after = text[ym.end() :].strip().lstrip(".").strip()
    after = re.sub(r"\s*https?://\S+\s*$", "", after).strip()
    after = re.sub(r"\s*doi:\s*[^\s]+\s*$", "", after, flags=re.I).strip()
    after = after.rstrip(".")
    sw = "[Computer software]" in after
    after_clean = after.replace("[Computer software]", "").strip().lstrip(".").strip()

    title, source = _split_title_and_source(after_clean)
    ref["title"] = _strip_surrounding_quotes(title).rstrip(" ,.;'\"“”‘’")
    journal, vol, issue, pages, is_book = _parse_source_fields(source)
    ref["journal"] = journal
    ref["vol"] = vol
    ref["issue"] = issue
    ref["pages"] = pages
    ref["is_book"] = is_book or (not vol and not pages and bool(journal))

    # If there was no identifiable source segment, treat the whole remainder as title/book.
    if not source and not ref["journal"]:
        ref["title"] = _strip_surrounding_quotes(after_clean).rstrip(" ,.;'\"“”‘’")
        ref["is_book"] = True

    if sw:
        ref["title"] = (ref["title"] + " [Computer software]").strip()
    ref["key"] = _citation_key(ref["authors"], ref["year"], ref["ysuf"])
    return ref


def parse_apa_reference(text: str) -> dict:
    # Backward-compatible alias used by verifier and existing scripts.
    return parse_author_date_reference(text)


def _strip_reference_lead_marker(text: str) -> Tuple[Optional[int], str]:
    s = _norm_space(text)
    m = re.match(r"^\s*[\[(]?\s*(\d{1,4})\s*[\])\.]?\s*(.*)$", s)
    if not m:
        return None, s
    num_txt = m.group(1)
    if not num_txt:
        return None, s
    try:
        num = int(num_txt)
    except ValueError:
        return None, s
    if num > 500:
        return None, s
    return num, (m.group(2) or "").strip()


def _authors_look_plausible(authors: Sequence[Tuple[str, str]]) -> bool:
    if not authors:
        return False
    lead = (authors[0][0] or "").strip()
    if not lead:
        return False
    if re.search(r"\bunknown\b", lead, flags=re.I):
        return False
    if re.search(r"[\[\]{}]", lead):
        return False
    if _norm_text(lead) in {"one", "two", "three", "four", "five"}:
        return False
    tokens = re.findall(r"[A-Za-z][A-Za-z'\-]*", lead)
    if not tokens:
        return False
    if len(tokens) > 4 and not re.search(r"\bet al\b", lead, flags=re.I):
        return False
    return True


def _title_looks_plausible(title: str) -> bool:
    t = _norm_space(title)
    if not t:
        return False
    if re.fullmatch(r"\d+(?:\([^)]+\))?(?::\d+(?:-\d+)?)?", t):
        return False
    if len(re.findall(r"[A-Za-z]", t)) < 4:
        return False
    return True


def _split_compact_title_journal(text: str) -> Optional[dict]:
    s = _norm_space(text).strip(" .;,")
    if not s:
        return None
    patterns = [
        re.compile(
            r"^(?P<title>.+?)\s+(?P<journal>[A-Z][A-Za-z0-9&,\- ]+?)\s+"
            r"(?P<year>(?:19|20)\d{2}[a-z]?)\s*;\s*(?P<vol>\d+)\((?P<issue>[^)]+)\)\s*:\s*(?P<pages>[0-9\-–]+)$"
        ),
        re.compile(
            r"^(?P<title>.+?)\s+(?P<journal>[A-Z][A-Za-z0-9&,\- ]+?)\s+"
            r"(?P<vol>\d+)\((?P<issue>[^)]+)\)\s*(?P<pages>[0-9\-–]+)$"
        ),
        re.compile(
            r"^(?P<title>.+?)\s+(?P<journal>[A-Z][A-Za-z0-9&,\- ]+?)\s+(?P<year>(?:19|20)\d{2}[a-z]?)$"
        ),
    ]
    for pat in patterns:
        m = pat.match(s)
        if m:
            return {
                "title": _norm_space(m.group("title")),
                "journal": _norm_space(m.group("journal")),
                "year": (m.groupdict().get("year") or "").strip(),
                "vol": (m.groupdict().get("vol") or "").strip(),
                "issue": (m.groupdict().get("issue") or "").strip(),
                "pages": (m.groupdict().get("pages") or "").strip(),
            }
    return None


def _clean_journal_field(journal: str) -> str:
    j = _norm_space(journal).strip(" ,.;")
    if not j:
        return ""
    j = re.sub(
        r"\s*(?:19|20)\d{2}[a-z]?\s*;\s*\d+\([^)]*\)\s*:\s*[0-9\-–]+$",
        "",
        j,
    ).strip(" ,.;")
    j = re.sub(
        r"\s*\d+\([^)]*\)\s*[: ]\s*[0-9\-–]+$",
        "",
        j,
    ).strip(" ,.;")
    return j


def _reference_core_score(ref: Optional[dict], *, require_num: bool = False) -> int:
    if not ref:
        return -1
    score = 0
    if _authors_look_plausible(ref.get("authors", [])):
        score += 3
    year = str(ref.get("year", "") or "").strip()
    if re.fullmatch(r"(?:19|20)\d{2}[a-z]?", year):
        score += 2
    elif year == "n.d.":
        score += 1
    title = (ref.get("title", "") or ref.get("title_part", "")).strip()
    if _title_looks_plausible(title):
        score += 2
    elif title:
        score -= 2
    journal = (ref.get("journal", "") or "").strip()
    if journal:
        score += 2
    elif ref.get("is_book"):
        score += 1
    if (ref.get("vol") or ref.get("pages")):
        score += 1
    if require_num and isinstance(ref.get("num"), int):
        score += 1
    return score


def _needs_hybrid_fallback(ref: Optional[dict], *, require_num: bool = False) -> bool:
    if not ref:
        return True
    if require_num and not isinstance(ref.get("num"), int):
        return True
    if not _authors_look_plausible(ref.get("authors", [])):
        return True
    title = (ref.get("title", "") or ref.get("title_part", "")).strip()
    if not _title_looks_plausible(title):
        return True
    year = str(ref.get("year", "") or "").strip()
    if year and not re.fullmatch(r"(?:19|20)\d{2}[a-z]?|n\.d\.", year):
        return True
    if not (ref.get("journal") or ref.get("is_book")):
        return True
    min_score = 7 if require_num else 6
    return _reference_core_score(ref, require_num=require_num) < min_score


def _fallback_hybrid_core(
    text: str,
    *,
    fallback_num: Optional[int] = None,
) -> dict:
    raw = _norm_space((text or "").replace("\u00A0", " "))
    detected_num, lead_stripped = _strip_reference_lead_marker(raw)
    num = detected_num if detected_num is not None else fallback_num

    doi = _extract_doi(raw)
    work = lead_stripped
    work = re.sub(r"https?://(?:dx\.)?doi\.org/[^\s]+", "", work, flags=re.I)
    work = re.sub(r"\bdoi:\s*[^\s]+", "", work, flags=re.I)
    work = re.sub(r"https?://[^\s]+", "", work, flags=re.I)
    work = _norm_space(work).strip(" .;,")
    if not doi:
        doi = _extract_doi(work)

    # Candidate 1: existing author-date parser.
    p_ad = parse_author_date_reference(work)
    cand_ad = {
        "raw": raw,
        "num": num,
        "authors": p_ad.get("authors", []),
        "year": p_ad.get("year", "") or "",
        "ysuf": p_ad.get("ysuf", "") or "",
        "title": p_ad.get("title", "") or "",
        "journal": p_ad.get("journal", "") or "",
        "vol": p_ad.get("vol", "") or "",
        "issue": p_ad.get("issue", "") or "",
        "pages": p_ad.get("pages", "") or "",
        "is_book": bool(p_ad.get("is_book")),
        "doi": p_ad.get("doi", "") or doi,
    }

    # Candidate 2: Vancouver parser after coercing a numeric lead.
    cand_van = None
    looks_vancouver_like = bool(re.search(r"(?:19|20)\d{2}[a-z]?\s*;\s*\d+", work))
    if looks_vancouver_like:
        parse_num = num if isinstance(num, int) and num > 0 else 1
        p_van = parse_vancouver_reference_line(f"{parse_num}. {work}", target_style="apa7")
        if p_van:
            cand_van = {
                "raw": raw,
                "num": num,
                "authors": p_van.get("authors", []),
                "year": p_van.get("year", "") or "",
                "ysuf": "",
                "title": p_van.get("title_part", "") or "",
                "journal": p_van.get("journal", "") or "",
                "vol": p_van.get("vol", "") or "",
                "issue": p_van.get("issue", "") or "",
                "pages": p_van.get("pages", "") or "",
                "is_book": False,
                "doi": p_van.get("doi", "") or doi,
            }

    # Candidate 3: direct heuristic extraction.
    h_year = ""
    h_ysuf = ""
    h_authors: List[Tuple[str, str]] = []
    h_title = ""
    h_source = ""
    compact_journal = ""
    compact_vol = ""
    compact_issue = ""
    compact_pages = ""
    before = ""
    after = work

    ym = re.search(r"\((?P<year>(?:19|20)\d{2})(?P<ysuf>[a-z]?)\)", work)
    if ym:
        h_year = ym.group("year")
        h_ysuf = ym.group("ysuf") or ""
        before = work[: ym.start()].strip().rstrip(".,;")
        after = work[ym.end() :].strip().lstrip(").,;:- ")
        h_authors = _parse_apa_authors(before) if before else []
    else:
        ym2 = re.search(r"(?<!\d)(?P<year>(?:19|20)\d{2})(?P<ysuf>[a-z]?)(?!\d)", work)
        if ym2:
            h_year = ym2.group("year")
            h_ysuf = ym2.group("ysuf") or ""
            before = work[: ym2.start()].strip().rstrip(".,;")
            after = work[ym2.end() :].strip().lstrip(").,;:- ")
            h_authors = _parse_apa_authors(before) if before else []

    if after:
        h_title, h_source = _split_title_and_source(after)

    a_part, t_part, s_part = _split_authors_title_vancouver(work)
    v_authors = _parse_vancouver_author_tokens(a_part)
    if not _authors_look_plausible(h_authors) and _authors_look_plausible(v_authors):
        h_authors = v_authors
    if not _authors_look_plausible(h_authors):
        m2 = re.search(
            r"([A-Z][A-Za-z'\-]+)\s+([A-Z])\.,\s*([A-Z][A-Za-z'\-]+)\s+([A-Z])\.",
            work,
        )
        if m2:
            h_authors = [
                (m2.group(1), f"{m2.group(2)}."),
                (m2.group(3), f"{m2.group(4)}."),
            ]
        else:
            m1 = re.search(r"([A-Z][A-Za-z'\-]+)\s+([A-Z])\.", work)
            if m1:
                h_authors = [(m1.group(1), f"{m1.group(2)}.")]

    if (_title_looks_plausible(t_part) and not _title_looks_plausible(h_title)) or (not h_title and t_part):
        h_title = t_part
    if not h_source and s_part:
        h_source = s_part

    compact = _split_compact_title_journal(h_source or h_title)
    if compact:
        compact_journal = compact.get("journal", "")
        compact_vol = compact.get("vol", "")
        compact_issue = compact.get("issue", "")
        compact_pages = _clean_pages(compact.get("pages", ""))
        if not h_year and compact.get("year"):
            m_y = re.match(r"((?:19|20)\d{2})([a-z]?)$", compact["year"])
            if m_y:
                h_year = m_y.group(1)
                h_ysuf = m_y.group(2) or ""
        if not _title_looks_plausible(h_title):
            h_title = compact.get("title", h_title)
        if not h_source or h_source == h_title:
            h_source = compact_journal

    if not _title_looks_plausible(h_title):
        chunks = [c.strip(" .;,") for c in re.split(r"\.\s+", after or work) if c.strip()]
        for chunk in chunks:
            if _title_looks_plausible(chunk):
                h_title = chunk
                break
        if chunks and not h_source:
            h_source = chunks[-1]

    v_source_info = _parse_vancouver_source(h_source or s_part or work)
    if not h_year and v_source_info.get("year"):
        yraw = v_source_info["year"]
        m_y = re.match(r"((?:19|20)\d{2})([a-z]?)$", yraw)
        if m_y:
            h_year = m_y.group(1)
            h_ysuf = m_y.group(2) or ""

    journal, vol, issue, pages, is_book = _parse_source_fields(h_source)
    if not journal and compact_journal:
        journal = compact_journal
    if not journal and v_source_info.get("journal"):
        journal = v_source_info.get("journal", "")
    if not vol and compact_vol:
        vol = compact_vol
    if not vol and v_source_info.get("vol"):
        vol = v_source_info.get("vol", "")
    if not issue and compact_issue:
        issue = compact_issue
    if not issue and v_source_info.get("issue"):
        issue = v_source_info.get("issue", "")
    if not pages and compact_pages:
        pages = compact_pages
    if not pages and v_source_info.get("pages"):
        pages = _clean_pages(v_source_info.get("pages", ""))

    if not journal:
        jm = re.search(
            r"(Journal of [A-Za-z0-9 ,&\-]+)\s+\d+\(",
            h_title or h_source or work,
        )
        if jm:
            journal = jm.group(1).strip()

    cand_heur = {
        "raw": raw,
        "num": num,
        "authors": h_authors,
        "year": h_year,
        "ysuf": h_ysuf,
        "title": _strip_surrounding_quotes(h_title).rstrip(" ,.;'\""),
        "journal": _clean_journal_field(journal),
        "vol": vol.strip(),
        "issue": issue.strip(),
        "pages": pages.strip(),
        "is_book": bool(is_book) or (not vol and not pages and bool(journal)),
        "doi": doi,
    }

    candidates = [cand_ad, cand_heur]
    if cand_van is not None:
        candidates.append(cand_van)
    best = max(candidates, key=_reference_core_score)

    best = dict(best)
    best["raw"] = raw
    best["num"] = num
    best["doi"] = best.get("doi", "") or doi
    best["title"] = _strip_surrounding_quotes((best.get("title", "") or "")).rstrip(" ,.;'\"")
    best["journal"] = _clean_journal_field(best.get("journal", ""))
    if not best.get("year"):
        best["year"] = "n.d."
        best["ysuf"] = ""
    return best


def _author_phrase_from_authors(authors: Sequence[Tuple[str, str]], style: str) -> str:
    phrase = _citation_key(authors, "2000", "").rsplit(", ", 1)[0]
    if style == "harvard":
        phrase = phrase.replace(" & ", " and ")
    return phrase


def _core_to_author_date_ref(core: dict) -> dict:
    year = str(core.get("year", "") or "").strip() or "n.d."
    ysuf = str(core.get("ysuf", "") or "").strip()
    if year == "n.d.":
        ysuf = ""
    ref = dict(
        raw=core.get("raw", ""),
        authors=core.get("authors", []),
        year=year,
        ysuf=ysuf,
        title=core.get("title", ""),
        journal=core.get("journal", ""),
        vol=core.get("vol", ""),
        issue=core.get("issue", ""),
        pages=core.get("pages", ""),
        key="",
        num=core.get("num"),
        is_book=bool(core.get("is_book")),
        is_nature=False,
        doi=core.get("doi", ""),
    )
    ref["key"] = _citation_key(ref["authors"], ref["year"], ref["ysuf"])
    return ref


def _core_to_vancouver_ref(core: dict, *, target_style: str, fallback_num: int) -> dict:
    num = core.get("num")
    if not isinstance(num, int) or num < 1:
        num = fallback_num
    year = str(core.get("year", "") or "").strip() or "n.d."
    ysuf = str(core.get("ysuf", "") or "").strip()
    if year != "n.d." and ysuf and not year.endswith(ysuf):
        year = f"{year}{ysuf}"
    authors = core.get("authors", [])
    author_phrase = _author_phrase_from_authors(authors, style=target_style)
    return {
        "num": num,
        "raw": core.get("raw", ""),
        "authors_part": _vancouver_authors(authors),
        "authors": authors,
        "title_part": core.get("title", ""),
        "source_part": core.get("journal", ""),
        "journal": core.get("journal", ""),
        "year": year,
        "ysuf": "",
        "vol": core.get("vol", ""),
        "issue": core.get("issue", ""),
        "pages": core.get("pages", ""),
        "doi": core.get("doi", ""),
        "author_phrase": author_phrase,
        "parenthetical": f"{author_phrase}, {year}",
    }


def _parse_author_date_reference_auto(text: str) -> Tuple[dict, bool]:
    parsed = parse_author_date_reference(text)
    if not _needs_hybrid_fallback(parsed):
        return parsed, False
    fallback = _core_to_author_date_ref(_fallback_hybrid_core(text))
    if _reference_core_score(fallback) >= _reference_core_score(parsed):
        return fallback, True
    return parsed, False


def _parse_vancouver_reference_auto(
    text: str,
    *,
    target_style: str,
    fallback_num: int,
) -> Tuple[Optional[dict], bool]:
    parsed = parse_vancouver_reference_line(text, target_style=target_style)
    if parsed and not _needs_hybrid_fallback(parsed, require_num=True):
        return parsed, False
    fallback = _core_to_vancouver_ref(
        _fallback_hybrid_core(text, fallback_num=fallback_num),
        target_style=target_style,
        fallback_num=fallback_num,
    )
    if parsed is None or _reference_core_score(fallback, require_num=True) >= _reference_core_score(parsed, require_num=True):
        return fallback, True
    return parsed, False


def _ensure_unique_reference_numbers(refs: Sequence[dict]) -> None:
    used: set[int] = set()
    next_num = 1
    for ref in refs:
        num = ref.get("num")
        if isinstance(num, int) and num > 0 and num not in used:
            used.add(num)
            next_num = max(next_num, num + 1)
            continue
        while next_num in used:
            next_num += 1
        ref["num"] = next_num
        used.add(next_num)
        next_num += 1


def _vancouver_authors(authors: Sequence[Tuple[str, str]]) -> str:
    parts: List[str] = []
    for i, (sn, ini) in enumerate(authors):
        if re.search(r"\bet al\b", sn, flags=re.I):
            parts.append("et al.")
            break
        if i >= 6:
            parts.append("et al.")
            break
        ci = ini.replace(".", "").replace(" ", "")
        parts.append(f"{sn} {ci}" if ci else sn)
    return ", ".join(parts)


def format_vancouver_reference(ref: dict, num: int) -> str:
    au = _vancouver_authors(ref["authors"]).strip().rstrip(".")
    ti = _strip_surrounding_quotes(ref["title"]).rstrip(" ,.;'\"“”‘’")
    ti_txt = _end_punct(ti)
    ja = JABBR.get(ref["journal"], ref["journal"])
    yr = ref["year"] + ref.get("ysuf", "")
    if ref.get("is_nature"):
        line = f"{num}. {au}. {ti_txt} {ja}. {yr};{ref['vol']}:{ref['pages']}."
        if ref.get("doi"):
            line = line.rstrip(".") + f". doi:{ref['doi']}."
        return line
    if ref.get("is_book"):
        line = f"{num}. {au}. {ti_txt} {ja}; {yr}."
        if ref.get("doi"):
            line = line.rstrip(".") + f". doi:{ref['doi']}."
        return line

    if ref.get("vol"):
        src = f"{ja}. {yr};{ref['vol']}"
        if ref["issue"]:
            src += f"({ref['issue']})"
        if ref["pages"]:
            src += f":{ref['pages']}"
    else:
        src = f"{ja}. {yr}"
        if ref["pages"]:
            src += f":{ref['pages']}"

    line = f"{num}. {au}. {ti_txt} {src}."
    if ref.get("doi"):
        line = line.rstrip(".") + f". doi:{ref['doi']}."
    return line


def format_numeric_ranges(numbers: Sequence[int]) -> str:
    if not numbers:
        return ""
    ns = sorted(set(numbers))
    ranges: List[Tuple[int, int]] = []
    s = ns[0]
    e = ns[0]
    for n in ns[1:]:
        if n == e + 1:
            e = n
        else:
            ranges.append((s, e))
            s = e = n
    ranges.append((s, e))
    parts: List[str] = []
    for a, b in ranges:
        parts.append(str(a) if a == b else (f"{a},{b}" if b - a == 1 else f"{a}-{b}"))
    return ",".join(parts)


class ApaToVancouverEngine:
    _ENTRY = re.compile(r"^\s*(.+?),\s*(\d{4}[a-z]?)(?:,\s*(.*))?\s*$")

    def __init__(self, refs: Sequence[dict]) -> None:
        self.refs = list(refs)
        self.by_key = {r["key"]: r for r in refs}
        self.order: List[str] = []
        self.key2num: Dict[str, int] = {}
        self._next = 1
        self._alt: Dict[str, str] = {}
        self._base_map: Dict[str, List[str]] = {}
        self.ambiguous_bases: List[str] = []
        self._narrative_pat = re.compile(r"$^")
        self._build_alt_keys()
        self._build_base_map()
        self._build_narrative_pattern()

    def _build_alt_keys(self) -> None:
        for key in self.by_key:
            k = _norm_space(key)
            self._alt[k] = key
            self._alt[_norm_space(k.replace(" & ", " and "))] = key

    def _build_base_map(self) -> None:
        base_map: Dict[str, List[str]] = {}
        for key in self.by_key:
            base = _norm_space(_citation_base_key(key))
            alt_base = _norm_space(base.replace(" & ", " and "))
            base_map.setdefault(base, []).append(key)
            if alt_base != base:
                base_map.setdefault(alt_base, []).append(key)
        for bk, vals in base_map.items():
            uniq = []
            for v in vals:
                if v not in uniq:
                    uniq.append(v)
            self._base_map[bk] = uniq

    def _build_narrative_pattern(self) -> None:
        author_parts: List[str] = []
        for key in self.by_key:
            if ", " in key:
                author_parts.append(key.rsplit(", ", 1)[0])
        variants = set()
        for ap in author_parts:
            variants.add(_norm_space(ap))
            variants.add(_norm_space(ap.replace(" & ", " and ")))
        if not variants:
            return
        alt = "|".join(sorted((re.escape(v) for v in variants), key=len, reverse=True))
        self._narrative_pat = re.compile(rf"(?<!\w)({alt})\s+\((\d{{4}}[a-z]?)([^)]*)\)")

    def _extract_key(self, entry: str) -> Optional[str]:
        m = self._ENTRY.match(_norm_space(entry))
        if not m:
            return None
        return f"{_norm_space(m.group(1))}, {m.group(2)}"

    def _extract_key_and_suffix(self, entry: str) -> Tuple[Optional[str], str]:
        m = self._ENTRY.match(_norm_space(entry))
        if not m:
            return None, ""
        key = f"{_norm_space(m.group(1))}, {m.group(2)}"
        suffix = (m.group(3) or "").strip()
        return key, suffix

    def _match_key(self, text: str) -> Optional[str]:
        t = _norm_space(text)
        c = self._alt.get(t)
        if c:
            return c
        t2 = _norm_space(t.replace(" and ", " & "))
        c2 = self._alt.get(t2)
        if c2:
            return c2

        base = _norm_space(_citation_base_key(t2))
        opts = self._base_map.get(base, [])
        if len(opts) == 1:
            return opts[0]
        if len(opts) > 1 and base not in self.ambiguous_bases:
            self.ambiguous_bases.append(base)
        return None

    def _record(self, key: str) -> Optional[int]:
        canonical = self._match_key(key)
        if not canonical:
            return None
        if canonical not in self.key2num:
            self.key2num[canonical] = self._next
            self.order.append(canonical)
            self._next += 1
        return self.key2num[canonical]

    def scan(self, paragraphs: Sequence) -> None:
        for para in paragraphs:
            t = full_text(para)
            events: List[Tuple[int, str, object]] = []
            for m in re.finditer(r"\(([^)]+)\)", t):
                events.append((m.start(), "paren", m))
            for m in self._narrative_pat.finditer(t):
                events.append((m.start(), "narr", m))
            events.sort(key=lambda x: x[0])

            for _pos, kind, m in events:
                if kind == "paren":
                    content = m.group(1)
                    content = re.sub(r"^(?:e\.g\.,|see|cf\.|i\.e\.,)\s*", "", content)
                    for part in content.split(";"):
                        key_try = self._extract_key(part)
                        if key_try:
                            self._record(key_try)
                else:
                    key_try = f"{_norm_space(m.group(1))}, {m.group(2)}"
                    self._record(key_try)

    def _scan_parenthetical(self, t: str) -> None:
        for m in re.finditer(r"\(([^)]+)\)", t):
            content = m.group(1)
            content = re.sub(r"^(?:e\.g\.,|see|cf\.|i\.e\.,)\s*", "", content)
            for part in content.split(";"):
                key_try = self._extract_key(part)
                if key_try:
                    self._record(key_try)

    def _scan_narrative(self, t: str) -> None:
        for m in self._narrative_pat.finditer(t):
            key_try = f"{_norm_space(m.group(1))}, {m.group(2)}"
            self._record(key_try)

    def replace(self, paragraphs: Sequence) -> Dict[str, int]:
        stats = {"paren": 0, "narr": 0}
        for para in paragraphs:
            self._replace_paragraph(para, stats)
        return stats

    def _replace_paragraph(self, para, stats: Dict[str, int]) -> None:
        t = full_text(para)
        paren_reps: List[Tuple[int, int, str]] = []
        for m in re.finditer(r"\(([^)]+)\)", t):
            content = m.group(1)
            pm = re.match(r"^((?:e\.g\.,|see|cf\.|i\.e\.,)\s*)", content)
            prefix = pm.group(1) if pm else ""
            cit_part = content[len(prefix) :]
            entries = [e.strip() for e in cit_part.split(";") if e.strip()]
            nums: List[int] = []
            rebuilt_entries: List[str] = []
            found_any = False
            all_entries_are_citations = True
            all_entries_are_pure = True
            for entry in entries:
                key_try, suffix = self._extract_key_and_suffix(entry)
                if key_try:
                    ck = self._match_key(key_try)
                    if ck and ck in self.key2num:
                        num = self.key2num[ck]
                        nums.append(num)
                        if suffix:
                            rebuilt_entries.append(f"[{num}], {suffix}")
                            all_entries_are_pure = False
                        else:
                            rebuilt_entries.append(f"[{num}]")
                        found_any = True
                    else:
                        all_entries_are_citations = False
                        rebuilt_entries.append(entry)
                else:
                    all_entries_are_citations = False
                    rebuilt_entries.append(entry)
            if found_any:
                tag = format_numeric_ranges(nums)
                if all_entries_are_citations and all_entries_are_pure:
                    replacement = f"({prefix}[{tag}])" if prefix else f"[{tag}]"
                else:
                    mixed = "; ".join(rebuilt_entries)
                    replacement = f"({prefix}{mixed})" if prefix else f"({mixed})"
                paren_reps.append((m.start(), m.end(), replacement))
                stats["paren"] += 1
        if paren_reps:
            replace_in_runs(para, paren_reps)

        t = full_text(para)
        reps: List[Tuple[int, int, str]] = []
        for m in self._narrative_pat.finditer(t):
            key_try = f"{_norm_space(m.group(1))}, {m.group(2)}"
            ck = self._match_key(key_try)
            if ck and ck in self.key2num:
                num = self.key2num[ck]
                suffix = (m.group(3) or "").strip()
                if suffix.startswith(","):
                    suffix = suffix[1:].strip()
                replacement = f"[{num}]"
                if suffix:
                    replacement += f" ({suffix})"
                year_start = m.start(2) - 1
                year_end = m.end(3) + 1
                reps.append((year_start, year_end, replacement))
                stats["narr"] += 1
        if reps:
            replace_in_runs(para, reps)


def _normalize_author_phrase(text: str) -> str:
    text = unicodedata.normalize("NFKD", text or "")
    text = text.encode("ascii", "ignore").decode("ascii")
    text = text.lower().replace(" and ", " & ")
    text = re.sub(r"[^a-z0-9& ]+", " ", text)
    return re.sub(r"\s+", " ", text).strip()


def _inside_parentheses(text: str, idx: int) -> bool:
    depth = 0
    for ch in text[:idx]:
        if ch == "(":
            depth += 1
        elif ch == ")" and depth > 0:
            depth -= 1
    return depth > 0


def _expand_num_token(token: str) -> List[int]:
    token = token.replace("–", "-")
    out: List[int] = []
    for part in token.split(","):
        part = part.strip()
        if not part:
            continue
        if "-" in part:
            a, b = part.split("-", 1)
            if a.isdigit() and b.isdigit():
                out.extend(range(int(a), int(b) + 1))
        elif part.isdigit():
            out.append(int(part))
    return out


def _surname_from_vancouver_author(token: str) -> str:
    tok = token.strip().rstrip(".")
    words = tok.split()
    while words and re.fullmatch(r"[A-Z]{1,4}", words[-1]):
        words.pop()
    if not words:
        return tok.split()[0] if tok.split() else tok
    return " ".join(words)


def _build_author_phrase_from_vancouver(authors_part: str, style: str = "apa7") -> str:
    s = authors_part.strip()
    if not s:
        return "Unknown"
    had_et_al = "et al" in s.lower()
    cleaned = s.replace("et al.", "").replace("et al", "")
    tokens = [t.strip() for t in cleaned.split(",") if t.strip()]
    surnames = [_surname_from_vancouver_author(t) for t in tokens]
    surnames = [s for s in surnames if s]
    if not surnames:
        return "Unknown"
    if had_et_al or len(surnames) >= 3:
        return f"{surnames[0]} et al."
    if len(surnames) == 2:
        conj = "&" if style == "apa7" else "and"
        return f"{surnames[0]} {conj} {surnames[1]}"
    return surnames[0]


def _split_authors_title_vancouver(rest: str) -> Tuple[str, str, str]:
    first_dot = rest.find(". ")
    if first_dot == -1:
        return rest.strip(), "", ""
    authors_part = rest[:first_dot].strip()
    rem = rest[first_dot + 2 :].strip()
    second_dot = rem.find(". ")
    if second_dot == -1:
        return authors_part, rem.rstrip("."), ""
    title_part = rem[:second_dot].strip().rstrip(".")
    source_part = rem[second_dot + 2 :].strip()
    return authors_part, title_part, source_part


def _parse_vancouver_author_tokens(authors_part: str) -> List[Tuple[str, str]]:
    s = authors_part.strip()
    if not s:
        return []
    had_et_al = "et al" in s.lower()
    cleaned = s.replace("et al.", "").replace("et al", "")
    tokens = [t.strip() for t in cleaned.split(",") if t.strip()]
    out: List[Tuple[str, str]] = []
    for token in tokens:
        words = token.split()
        if not words:
            continue
        surname = _surname_from_vancouver_author(token)
        ini_words = words[len(surname.split()) :]
        initials = "".join(ch for ch in "".join(ini_words) if ch.isalpha()).upper()
        ini_fmt = " ".join(f"{ch}." for ch in initials)
        out.append((surname, ini_fmt))
    if had_et_al and out:
        out.append(("et al.", ""))
    return out


def _parse_vancouver_source(source: str) -> dict:
    src = source.strip().rstrip(".")
    year = ""
    vol = ""
    issue = ""
    pages = ""
    journal = ""
    doi = _extract_doi(src)
    src = re.sub(r"https?://(?:dx\.)?doi\.org/[^\s]+", "", src, flags=re.I)
    src = re.sub(r"\bdoi:\s*[^\s]+", "", src, flags=re.I)
    src = re.sub(r"https?://[^\s]+", "", src, flags=re.I)
    src = _norm_space(src).rstrip(".")

    ym = re.search(r"(?<!\d)((?:19|20)\d{2}[a-z]?)(?!\d)", src)
    if ym:
        year = ym.group(1)
        journal = src[: ym.start()].rstrip(" .;")
        tail = src[ym.end() :].strip()
        vim = re.search(r"^;\s*(\d+)\s*(?:\(([^)]+)\))?\s*(?::\s*([^\s]+(?:\s*[^\s]+)*))?", tail)
        if vim:
            vol = vim.group(1) or ""
            issue = (vim.group(2) or "").strip()
            pages = (vim.group(3) or "").strip().rstrip(".")
    else:
        journal = src

    if not pages:
        pm = re.search(r"\bpp?\.\s*([0-9]+(?:\s*[-–]\s*[0-9]+)?)", src, flags=re.I)
        if pm:
            pages = pm.group(1).replace(" ", "")

    return {
        "journal": journal.strip().rstrip(","),
        "year": year,
        "vol": vol,
        "issue": issue,
        "pages": pages,
        "doi": doi,
    }


def parse_vancouver_reference_line(line: str, target_style: str = "apa7") -> Optional[dict]:
    m = re.match(r"^(\d+)\.\s+(.*)$", line.strip())
    if not m:
        return None
    num = int(m.group(1))
    rest = m.group(2).strip()
    authors_part, title_part, source_part = _split_authors_title_vancouver(rest)
    source_info = _parse_vancouver_source(source_part)
    year = source_info["year"] or "n.d."
    authors = _parse_vancouver_author_tokens(authors_part)
    author_phrase = _build_author_phrase_from_vancouver(authors_part, style=target_style)
    return {
        "num": num,
        "raw": rest,
        "authors_part": authors_part,
        "authors": authors,
        "title_part": title_part,
        "source_part": source_part,
        "journal": source_info["journal"],
        "year": year,
        "vol": source_info["vol"],
        "issue": source_info["issue"],
        "pages": source_info["pages"],
        "doi": source_info["doi"],
        "author_phrase": author_phrase,
        "parenthetical": f"{author_phrase}, {year}",
    }


def _format_author_date_authors(authors: Sequence[Tuple[str, str]], style: str) -> str:
    if not authors:
        return "Unknown"
    if len(authors) == 1 and re.search(r"\bet al\b", authors[0][0], flags=re.I):
        return "et al."

    parts: List[str] = []
    for sn, ini in authors:
        if re.search(r"\bet al\b", sn, flags=re.I):
            if parts:
                return f"{parts[0]} et al."
            return "et al."
        if style == "harvard":
            ini_text = (ini or "").replace(" ", "")
        else:
            ini_text = ini or ""
        parts.append(f"{sn}, {ini_text}".rstrip(", ").strip())

    if not parts:
        return "Unknown"
    if len(parts) == 1:
        return parts[0]
    if style == "harvard":
        if len(parts) == 2:
            return f"{parts[0]} and {parts[1]}"
        return ", ".join(parts[:-1]) + f" and {parts[-1]}"
    if len(parts) == 2:
        return f"{parts[0]}, & {parts[1]}"
    return ", ".join(parts[:-1]) + f", & {parts[-1]}"


def format_author_date_reference(ref: dict, style: str = "apa7") -> str:
    authors = _format_author_date_authors(ref.get("authors", []), style=style)
    year = (ref.get("year", "") + ref.get("ysuf", "")).strip() or "n.d."
    title = (ref.get("title", "") or ref.get("title_part", "")).strip().rstrip(".")
    journal = (ref.get("journal", "") or "").strip().rstrip(".")
    vol = (ref.get("vol", "") or "").strip()
    issue = (ref.get("issue", "") or "").strip()
    pages = (ref.get("pages", "") or "").strip()
    doi = (ref.get("doi", "") or "").strip()

    if style == "harvard":
        title_fmt = f"'{title}'" if title else ""
        line = f"{authors} ({year}) {title_fmt}".strip()
        source_bits: List[str] = []
        if journal:
            source_bits.append(journal)
        if vol:
            if issue:
                source_bits.append(f"{vol}({issue})")
            else:
                source_bits.append(vol)
        if pages:
            pp = pages
            if not re.match(r"^(p|pp)\.", pages, flags=re.I):
                pp = f"pp. {pages}"
            source_bits.append(pp)
        if source_bits:
            line += ", " + ", ".join(source_bits)
        line = line.rstrip(", ").rstrip(".") + "."
        if doi:
            line += f" doi:{doi}."
        return line

    # APA7 best effort
    line = f"{authors} ({year})."
    if title:
        line += f" {title}."
    if journal:
        line += f" {journal}"
    if vol:
        if issue:
            line += f", {vol}({issue})"
        else:
            line += f", {vol}"
    if pages:
        line += f", {pages}"
    line = line.rstrip(", ").rstrip(".") + "."
    if doi:
        line += f" doi:{doi}."
    return line


def _looks_narrative_context(text: str, pos: int, author_phrase: str) -> bool:
    left = text[:pos]
    m = re.search(
        r"([A-Za-z][A-Za-z'`\-]*(?:\s+[A-Za-z][A-Za-z'`\-]*)?(?:\s+(?:&|and)\s+[A-Za-z][A-Za-z'`\-]*)?(?:\s+et\s+al\.)?)\s*$",
        left,
    )
    if not m:
        return False
    lhs = _normalize_author_phrase(m.group(1))
    rhs = _normalize_author_phrase(author_phrase)
    return lhs == rhs


def convert_author_date_to_vancouver(
    doc: Document,
    keep_uncited: bool = True,
    *,
    ref_header_n: int = 1,
) -> Dict[str, int]:
    ref_idx = ensure_reference_header(doc, occurrence=ref_header_n)
    if ref_idx == -1:
        raise RuntimeError(
            "Could not find reference header or infer a reference list automatically."
        )
    _set_references_header(doc, ref_idx)
    ref_paras = _collect_ref_paragraphs(doc, ref_idx)
    if not ref_paras:
        raise RuntimeError("No references found below reference header.")
    refs: List[dict] = []
    hybrid_count = 0
    for _, txt in ref_paras:
        parsed, used_hybrid = _parse_author_date_reference_auto(txt)
        refs.append(parsed)
        if used_hybrid:
            hybrid_count += 1
    refs, duplicate_count = _collapse_exact_duplicate_references(refs)
    _assign_year_suffixes(refs)

    engine = ApaToVancouverEngine(refs)
    body_paras = collect_body_paragraphs_before_reference(doc, ref_idx)
    hf_paras = collect_header_footer_paragraphs(doc)
    scan_paras = body_paras + hf_paras
    engine.scan(scan_paras)
    if engine.ambiguous_bases:
        raise RuntimeError(
            "Ambiguous same-author/year citations found without disambiguating suffixes (a/b). "
            "Please update in-text citations to include year suffixes for: "
            + ", ".join(engine.ambiguous_bases[:8])
        )
    stats = engine.replace(scan_paras)

    ordered_refs: List[dict] = []
    for key in engine.order:
        r = engine.by_key.get(key)
        if r:
            r["num"] = engine.key2num[key]
            ordered_refs.append(r)
    if keep_uncited:
        for r in refs:
            if r["key"] not in engine.key2num:
                r["num"] = engine._next
                engine._next += 1
                ordered_refs.append(r)

    _remove_paragraphs_by_indices(doc, [idx for idx, _ in ref_paras])
    new_ref_idx = _find_ref_header(doc)
    _set_references_header(doc, new_ref_idx)
    lines = [format_vancouver_reference(r, r["num"]) for r in ordered_refs]
    _insert_reference_lines(doc, new_ref_idx, lines)

    return {
        "body_unique_citations": len(engine.order),
        "paren_replacements": stats["paren"],
        "narr_replacements": stats["narr"],
        "reference_count": len(lines),
        "header_footer_paragraphs": len(hf_paras),
        "hybrid_normalized": hybrid_count,
        "duplicate_refs_collapsed": duplicate_count,
    }


def convert_apa7_to_vancouver(
    doc: Document,
    keep_uncited: bool = True,
    *,
    ref_header_n: int = 1,
) -> Dict[str, int]:
    # Backward-compatible name.
    return convert_author_date_to_vancouver(doc, keep_uncited=keep_uncited, ref_header_n=ref_header_n)


def convert_harvard_to_vancouver(
    doc: Document,
    keep_uncited: bool = True,
    *,
    ref_header_n: int = 1,
) -> Dict[str, int]:
    return convert_author_date_to_vancouver(doc, keep_uncited=keep_uncited, ref_header_n=ref_header_n)


def convert_vancouver_to_author_date(
    doc: Document,
    target_style: str = "apa7",
    sort_references: bool = True,
    *,
    ref_header_n: int = 1,
) -> Dict[str, int]:
    ref_idx = ensure_reference_header(doc, occurrence=ref_header_n)
    if ref_idx == -1:
        raise RuntimeError(
            "Could not find reference header or infer a reference list automatically."
        )
    _set_references_header(doc, ref_idx)
    ref_paras = _collect_ref_paragraphs(doc, ref_idx)
    if not ref_paras:
        raise RuntimeError("No references found below reference header.")

    target_style = target_style.lower().strip()
    if target_style not in {"apa7", "harvard"}:
        raise RuntimeError("target_style must be 'apa7' or 'harvard'.")

    parsed: List[dict] = []
    hybrid_count = 0
    for i, (_, txt) in enumerate(ref_paras, start=1):
        p, used_hybrid = _parse_vancouver_reference_auto(
            txt,
            target_style=target_style,
            fallback_num=i,
        )
        if p is None:
            continue
        parsed.append(p)
        if used_hybrid:
            hybrid_count += 1
    _ensure_unique_reference_numbers(parsed)
    _assign_year_suffixes(parsed)
    for p in parsed:
        p["parenthetical"] = f"{p['author_phrase']}, {p['year']}{p.get('ysuf', '')}"
    by_num = {p["num"]: p for p in parsed}
    body_paras = collect_body_paragraphs_before_reference(doc, ref_idx)
    hf_paras = collect_header_footer_paragraphs(doc)
    work_paras = body_paras + hf_paras

    repl_count = 0
    for para in work_paras:
        t = full_text(para)
        reps: List[Tuple[int, int, str]] = []
        for m in re.finditer(r"\[(\d+(?:[,\-–]\d+)*)\]", t):
            nums = _expand_num_token(m.group(1))
            if not nums:
                continue
            labels: List[str] = []
            for n in nums:
                if n in by_num:
                    labels.append(by_num[n]["parenthetical"])
            if not labels:
                continue
            dedup: List[str] = []
            for lb in labels:
                if lb not in dedup:
                    dedup.append(lb)
            inside = _inside_parentheses(t, m.start())
            if len(nums) == 1 and nums[0] in by_num:
                one = by_num[nums[0]]
                if _looks_narrative_context(t, m.start(), one["author_phrase"]):
                    replacement = one["year"] if inside else f"({one['year']})"
                else:
                    cite_text = "; ".join(dedup)
                    replacement = cite_text if inside else f"({cite_text})"
            else:
                cite_text = "; ".join(dedup)
                replacement = cite_text if inside else f"({cite_text})"
            reps.append((m.start(), m.end(), replacement))
        if reps:
            replace_in_runs(para, reps)
            repl_count += len(reps)

    _remove_paragraphs_by_indices(doc, [idx for idx, _ in ref_paras])
    new_ref_idx = _find_ref_header(doc)
    _set_references_header(doc, new_ref_idx)
    refs_out = [format_author_date_reference(r, style=target_style) for r in parsed]
    if sort_references:
        def _sort_key(ref_line: str) -> Tuple[str, str, str]:
            p = parse_author_date_reference(ref_line) or {}
            return (
                _normalize_author_phrase(p.get("key", "").split(", ", 1)[0] if p.get("key") else ""),
                p.get("year", ""),
                _norm_space(ref_line.lower()),
            )
        refs_out = sorted(refs_out, key=_sort_key)
    _insert_reference_lines(doc, new_ref_idx, refs_out)

    return {
        "citation_replacements": repl_count,
        "reference_count": len(refs_out),
        "header_footer_paragraphs": len(hf_paras),
        "hybrid_normalized": hybrid_count,
    }


def convert_vancouver_to_apa7(
    doc: Document,
    sort_references: bool = True,
    *,
    ref_header_n: int = 1,
) -> Dict[str, int]:
    return convert_vancouver_to_author_date(
        doc,
        target_style="apa7",
        sort_references=sort_references,
        ref_header_n=ref_header_n,
    )


def convert_vancouver_to_harvard(
    doc: Document,
    sort_references: bool = True,
    *,
    ref_header_n: int = 1,
) -> Dict[str, int]:
    return convert_vancouver_to_author_date(
        doc,
        target_style="harvard",
        sort_references=sort_references,
        ref_header_n=ref_header_n,
    )


def _restyle_parenthetical_author(author_part: str, target_style: str) -> str:
    s = _norm_space(author_part)
    if target_style == "apa7":
        return re.sub(r"\s+and\s+", " & ", s, flags=re.I)
    return s.replace(" & ", " and ")


def convert_author_date_to_author_date(
    doc: Document,
    target_style: str = "apa7",
    *,
    ref_header_n: int = 1,
) -> Dict[str, int]:
    ref_idx = ensure_reference_header(doc, occurrence=ref_header_n)
    if ref_idx == -1:
        raise RuntimeError(
            "Could not find reference header or infer a reference list automatically."
        )
    _set_references_header(doc, ref_idx)
    ref_paras = _collect_ref_paragraphs(doc, ref_idx)
    if not ref_paras:
        raise RuntimeError("No references found below reference header.")
    refs: List[dict] = []
    hybrid_count = 0
    for _, txt in ref_paras:
        parsed, used_hybrid = _parse_author_date_reference_auto(txt)
        refs.append(parsed)
        if used_hybrid:
            hybrid_count += 1
    refs, duplicate_count = _collapse_exact_duplicate_references(refs)
    _assign_year_suffixes(refs)

    body_paras = collect_body_paragraphs_before_reference(doc, ref_idx)
    hf_paras = collect_header_footer_paragraphs(doc)
    work_paras = body_paras + hf_paras
    citation_restyles = 0
    entry_pat = re.compile(r"^\s*(.+?),\s*(\d{4}[a-z]?)\s*$")
    for para in work_paras:
        t = full_text(para)
        reps: List[Tuple[int, int, str]] = []
        for m in re.finditer(r"\(([^)]+)\)", t):
            content = m.group(1)
            pm = re.match(r"^((?:e\.g\.,|see|cf\.|i\.e\.,)\s*)", content)
            prefix = pm.group(1) if pm else ""
            cit_part = content[len(prefix) :]
            entries = [e.strip() for e in cit_part.split(";") if e.strip()]
            rebuilt: List[str] = []
            changed = False
            for entry in entries:
                em = entry_pat.match(_norm_space(entry))
                if not em:
                    rebuilt.append(entry)
                    continue
                new_author = _restyle_parenthetical_author(em.group(1), target_style=target_style)
                rebuilt.append(f"{new_author}, {em.group(2)}")
                changed = changed or (new_author != _norm_space(em.group(1)))
            if changed:
                replacement = f"({prefix}{'; '.join(rebuilt)})"
                reps.append((m.start(), m.end(), replacement))
        if reps:
            replace_in_runs(para, reps)
            citation_restyles += len(reps)

    _remove_paragraphs_by_indices(doc, [idx for idx, _ in ref_paras])
    new_ref_idx = _find_ref_header(doc)
    _set_references_header(doc, new_ref_idx)
    refs_out = [format_author_date_reference(r, style=target_style) for r in refs]
    _insert_reference_lines(doc, new_ref_idx, refs_out)

    return {
        "citation_restyles": citation_restyles,
        "reference_count": len(refs_out),
        "header_footer_paragraphs": len(hf_paras),
        "hybrid_normalized": hybrid_count,
        "duplicate_refs_collapsed": duplicate_count,
    }


def convert_apa7_to_harvard(doc: Document, *, ref_header_n: int = 1) -> Dict[str, int]:
    return convert_author_date_to_author_date(doc, target_style="harvard", ref_header_n=ref_header_n)


def convert_harvard_to_apa7(doc: Document, *, ref_header_n: int = 1) -> Dict[str, int]:
    return convert_author_date_to_author_date(doc, target_style="apa7", ref_header_n=ref_header_n)


def _normalize_mode(mode: str) -> str:
    m = (mode or "").strip().lower().replace("_", "-").replace(" ", "")
    aliases = {
        "apa7-to-vancouver": "a2v",
        "apa-to-vancouver": "a2v",
        "harvard-to-vancouver": "h2v",
        "harvard-vancouver": "h2v",
        "h2v": "h2v",
        "apa7-vancouver": "a2v",
        "apa2van": "a2v",
        "apa-to-van": "a2v",
        "a2v": "a2v",
        "vancouver-to-apa7": "v2a",
        "van-to-apa7": "v2a",
        "van-apa7": "v2a",
        "vancouver-to-harvard": "v2h",
        "van-to-harvard": "v2h",
        "van-harvard": "v2h",
        "v2h": "v2h",
        "van2apa": "v2a",
        "v2a": "v2a",
        "apa7-to-harvard": "a2h",
        "apa-to-harvard": "a2h",
        "a2h": "a2h",
        "harvard-to-apa7": "h2a",
        "harvard-to-apa": "h2a",
        "h2a": "h2a",
    }
    if m in aliases:
        return aliases[m]
    raise ValueError(f"Unsupported mode: {mode}")


def _default_output(input_path: Path, mode: str) -> Path:
    if mode in {"a2v", "h2v"}:
        suffix = "_vancouver.docx"
    elif mode in {"v2a", "h2a"}:
        suffix = "_apa7.docx"
    else:
        suffix = "_harvard.docx"
    return input_path.with_name(f"{input_path.stem}{suffix}")


def main() -> int:
    parser = argparse.ArgumentParser(
        description="Convert .docx citation/reference style between APA7, Harvard, and Vancouver."
    )
    parser.add_argument(
        "--mode",
        required=True,
        help=(
            "apa7-to-vancouver | harvard-to-vancouver | vancouver-to-apa7 | "
            "vancouver-to-harvard | apa7-to-harvard | harvard-to-apa7 (aliases supported)"
        ),
    )
    parser.add_argument("--input", required=True, help="Input .docx file path")
    parser.add_argument("--output", help="Output .docx file path")
    parser.add_argument(
        "--ref-header-n",
        type=int,
        default=1,
        help="1-based reference-header occurrence to use when multiple 'References' headers exist",
    )
    parser.add_argument(
        "--allow-field-codes",
        action="store_true",
        help="Allow conversion even when Zotero/EndNote/Mendeley field codes are detected",
    )
    parser.add_argument(
        "--allow-unsupported-parts",
        action="store_true",
        help="Allow conversion when citation-like text is detected in unsupported DOCX regions (text boxes, footnotes/endnotes)",
    )
    parser.add_argument(
        "--drop-uncited",
        action="store_true",
        help="For author-date->Vancouver only: omit uncited references from output list",
    )
    parser.add_argument(
        "--no-sort-apa",
        action="store_true",
        help="For Vancouver->author-date only: keep reference order instead of sorting",
    )
    args = parser.parse_args()

    mode = _normalize_mode(args.mode)
    in_path = Path(args.input)
    if not in_path.exists():
        print(f"ERROR: Input file not found: {in_path}")
        return 1
    out_path = Path(args.output) if args.output else _default_output(in_path, mode)

    print(f"Reading: {in_path}")
    doc = Document(str(in_path))
    preflight = preflight_docx(
        doc,
        ref_header_n=args.ref_header_n,
        allow_field_codes=args.allow_field_codes,
        allow_unsupported_parts=args.allow_unsupported_parts,
    )
    for w in preflight["warnings"]:
        print(f"[WARN] {w}")
    if preflight["failures"]:
        for f in preflight["failures"]:
            print(f"[ERROR] {f}")
        return 2

    try:
        if mode in {"a2v", "h2v"}:
            stats = convert_author_date_to_vancouver(
                doc,
                keep_uncited=not args.drop_uncited,
                ref_header_n=args.ref_header_n,
            )
            source_label = "APA7" if mode == "a2v" else "Harvard"
            print(f"Mode: {source_label} -> Vancouver")
            print(f"- unique citations in body: {stats['body_unique_citations']}")
            print(f"- parenthetical replacements: {stats['paren_replacements']}")
            print(f"- narrative replacements: {stats['narr_replacements']}")
            print(f"- header/footer paragraphs scanned: {stats['header_footer_paragraphs']}")
            print(f"- references written: {stats['reference_count']}")
        elif mode in {"v2a", "v2h"}:
            target_style = "apa7" if mode == "v2a" else "harvard"
            stats = convert_vancouver_to_author_date(
                doc,
                target_style=target_style,
                sort_references=not args.no_sort_apa,
                ref_header_n=args.ref_header_n,
            )
            print(f"Mode: Vancouver -> {target_style.upper()} (best effort)")
            print(f"- citation replacements: {stats['citation_replacements']}")
            print(f"- header/footer paragraphs scanned: {stats['header_footer_paragraphs']}")
            print(f"- references written: {stats['reference_count']}")
        else:
            target_style = "harvard" if mode == "a2h" else "apa7"
            stats = convert_author_date_to_author_date(
                doc,
                target_style=target_style,
                ref_header_n=args.ref_header_n,
            )
            source_label = "APA7" if mode == "a2h" else "Harvard"
            print(f"Mode: {source_label} -> {target_style.upper()} (best effort)")
            print(f"- citation restyles: {stats['citation_restyles']}")
            print(f"- header/footer paragraphs scanned: {stats['header_footer_paragraphs']}")
            print(f"- references written: {stats['reference_count']}")
    except RuntimeError as exc:
        print(f"ERROR: {exc}")
        return 1

    doc.save(str(out_path))
    print(f"Saved: {out_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
